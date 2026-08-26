"""
Single-Shot 批改器 —— 重新设计的核心架构
============================================
核心理念：将全文一次性交给大模型，由模型返回结构化 JSON（含所有批注+总评），
然后代码逐条调用 tools.py 的成熟工具函数注入到 Word 文档。

1. ✅ 总评不会缺失 —— JSON 结构强制包含，不存在"上下文裁剪丢掉"
2. ✅ 红绿不混杂 —— 一次性决定，不存在反复调用覆盖
3. ✅ 省 token —— 1 次 API 调用 vs 原来 10-40 次
4. ✅ 结构清晰 —— 段落定位由代码完成，模型只需描述问题
5. ✅ 错误隔离 —— LLM 只负责判断，不负责工具调用顺序
"""

import json
import os
import re
import shutil
import sys
import time
from pathlib import Path
from typing import Optional

from openai import OpenAI
from docx import Document

# 确保项目根目录在路径中
_project_root = str(Path(__file__).resolve().parent.parent)
if _project_root not in sys.path:
    sys.path.insert(0, _project_root)

from word_agent.config import (
    ARK_API_KEY, DEFAULT_MODEL, DEFAULT_BASE_URL,
    TEMPERATURE, REASONING_EFFORT, MAX_TOKENS,
    COMMENT_DIMENSIONS, SCORE_TOTAL,
    INTERVIEW_DIMENSIONS, INTERVIEW_SCORE_TOTAL,
)
from word_agent.tools import (
    tool_read_answer_pdf,
    tool_read_homework,
    tool_add_green_sentence,
    tool_add_red_sentence,
    tool_add_paragraph_comment,
    tool_add_strikethrough,
    tool_add_table_comment,
    tool_append_final_review,
    normalize_document_runs,
    cleanup_temp_files,
    _is_protected_para,
)


# ============================================================
# 通用批改评分标准（内联，不再依赖外部文件）
# ============================================================
# 来源：老师批改风格提炼、考研阅卷标准
# 设计原则：不绑定任何具体作业材料，适用于该类型的所有作业

# ---- 新闻评论（50 分制） ----
COMMENT_ANCHORS = {
    "立意": {
        "low": "无明确总论点，或只讲事不讲评（像议论文），或立意跑偏",
        "mid": "立意正确但偏泛/偏口号，或'评'的力度不足，或总论点位置/表述欠佳",
        "high": "总论点鲜明、有'评'（有态度有观点），开头点明并与标题呼应，观点有深度或新意",
    },
    "标题": {
        "low": "标题无观点、只描述事件，或明显不符评论标题规范",
        "mid": "标题切题但观点性不足，或偏学术/偏口号，或格式欠佳",
        "high": "观点鲜明、切中核心，有'评'，符合评论标题规范且有吸引力",
    },
    "论证逻辑": {
        "low": "论证无逻辑、东拉西扯，或纯罗列事实不评论，或无依据主观臆断",
        "mid": "有论证但部分段落只讲事缺观点收束，或论证较浅、论据单薄",
        "high": "每段做到'事实/事件→过渡→观点收束'闭环，层层递进、论据充分、无跳跃",
    },
    "行文结构": {
        "low": "结构混乱、段落职责不清，或缺少明确的总分结构",
        "mid": "结构基本完整但某环节缺失（缺分论点/开头只讲事/结尾空泛/引语不必要）",
        "high": "开头(由头+过渡+总立意)→主体(分论点围绕总论点)→结尾(升华呼应)完整清晰",
    },
    "语言表达": {
        "low": "语病较多、语言随意、不符评论语体，或材料人名/数据写错",
        "mid": "少量语病/口语化/网络语/标点问题，不影响理解",
        "high": "客观专业、符合评论语体，无网络语/口语化，标点用词规范",
    },
}
COMMENT_STRICTNESS = [
    "严格度：考研阅卷水平，宁严勿松；宁可给低分并说明原因，不可为鼓励而虚高。",
    "硬性校准：某维度只要符合'低'档描述，该维度必须落在 0-4 分；其他维度的优点不能抵消或抬高该维度。",
    "只讲事不讲评（像议论文、缺观点收束）= 立意 0-4 分，论证逻辑 ≤6 分。",
    "材料人名/数据/事实写错 = 语言表达与论证逻辑各扣 1-2 分（态度问题）。",
    "无依据的主观臆断 = 论证逻辑扣分，并用红色评语指出。",
    "引语/开头不必要且无效果 = 行文结构扣分，建议删掉。",
    "分数与问题联动：只要诊断中指出需要结构性修改，总分应落在 28-36 区间；仅当无此类实质问题才可给 40+。",
    "法无定法：立意好、论证漂亮、有新意，即使不按标准答案也应给 8-10 分。",
    "拉开差距：好文章敢给 40+，有硬伤敢给 ≤32，不要全部挤在中间。",
]
COMMENT_EXEMPLARS_TEXT = (
    "\n## 严格度示范（匿名样例，仅用于校准评分尺度）\n"
    "- 学生表现：正文分三段分别从'商家、博主、网民'展开，全是事实罗列，每段讲完事就结束，没有落到观点上；总论点模糊，整体更像议论文。  →  老师打分：立意4/标题5/论证逻辑6/行文结构7/语言表达7\n"
    "- 学生表现：总论点明确、标题切题，但第二个分论点只是现象描述、没有围绕总论点展开，且有两段'讲完事实就戛然而止'、缺观点收束。  →  老师打分：立意6/标题6/论证逻辑6/行文结构7/语言表达7\n"
    "- 学生表现：总论点鲜明、开头点明并与标题呼应，三个分论点层层递进、都围绕总论点，每段完成'事实→观点收束'，结尾升华有力。  →  老师打分：立意8/标题8/论证逻辑7/行文结构8/语言表达6"
)
# 评语写作风格（评论/消息共用）——以老师真实批改语气为基准，克制 AI 味
RUBRIC_STYLE_TEXT = (
    "\n## 评语风格\n"
    "- 绿色评语简短自然，但要点出好在哪：可以的，数据用得到位。/ 这句很自然，转折不生硬。\n"
    "- 严禁承接上文、总结分论点、继续叙述、分析论证、收束得当、点明坚持、结构完整、论证有力 这类套话\n"
    "- 红色纠错**不受字数限制**：问题句必须展开讲透——指出问题 + 给具体改法 + 必要时补一句原因/影响，如：弥漫与恐慌搭配不当，建议改为制造恐慌，这样语义更准确。\n"
    "- **禁止空话**：评语必须具体，落到原文的用词、结构、逻辑或材料事实上；严禁「有待加强」「继续努力」「整体不错」这类无信息量的套话。\n"
    "- 段落级评语只在整段有突出的结构/逻辑/立意问题或亮点时用，不要每段都加。\n"
)
COMMENT_RUBRIC_TEXT = "\n## 评分锚点（每维度满分 10，高=8-10 中=5-7 低=0-4）\n" + "\n".join(
    f"- 【{dim}】低：{levels['low']}；中：{levels['mid']}；高：{levels['high']}"
    for dim, levels in COMMENT_ANCHORS.items()
) + "\n\n## 严格度校准\n" + "\n".join(f"- {r}" for r in COMMENT_STRICTNESS) + RUBRIC_STYLE_TEXT + "\n" + COMMENT_EXEMPLARS_TEXT

# ---- 消息改写（40 分制） ----
XIAOXI_ANCHORS = {
    "标题": {
        "low": "标题与内容不符、核心事实缺失，或明显不符合消息标题规范",
        "mid": "标题基本准确但欠精炼，或未突出最新核心事实",
        "high": "复合标题，准确统摄全文、直击核心事实，精炼有吸引力",
    },
    "导语": {
        "low": "电头格式错误，或核心事实偏/缺失，或导语混杂大量背景/专家意见",
        "mid": "电头或要素有缺漏，核心事件位置欠佳，仍有专家意见等非核心内容",
        "high": "电头格式正确，4W 要素齐全，最新核心事件置顶（倒金字塔）",
    },
    "逻辑清晰、事实清楚": {
        "low": "结构混乱、非倒金字塔，或关键数据错误、事实张冠李戴",
        "mid": "结构基本清晰但个别段落顺序欠佳，或数据有个别疏漏",
        "high": "倒金字塔结构清晰，数据准确无遗漏，主体对导语逐层展开",
    },
    "语言通顺、紧扣材料": {
        "low": "语病较多、主观色彩浓，或大量照搬材料、明显超/欠字数",
        "mid": "语言基本通顺但个别口语化/冗余，或对材料运用不充分",
        "high": "客观中立、新闻书面语，紧扣材料、字数达标",
    },
}
XIAOXI_STRICTNESS = [
    "严格度：考研阅卷水平，宁严勿松；宁可给低分并说明原因，不可为鼓励而虚高。",
    "硬性校准：某维度只要符合'低'档描述，该维度必须落在 0-4 分；其他维度优点不能抵消。",
    "电头格式错误、核心事实偏差 = 导语大扣分（≤4/10）。",
    "数据写错、人名/机构简称不规范 = 逻辑与事实维度扣分，并红评指出。",
    "法无定法：结构创新、语言漂亮，即使不按标准答案也应给分。",
    "拉开差距：好稿敢给 34+，有硬伤敢给 ≤24。",
]
XIAOXI_RUBRIC_TEXT = "\n## 评分锚点（每维度满分 10，高=8-10 中=5-7 低=0-4）\n" + "\n".join(
    f"- 【{dim}】低：{levels['low']}；中：{levels['mid']}；高：{levels['high']}"
    for dim, levels in XIAOXI_ANCHORS.items()
) + "\n\n## 严格度校准\n" + "\n".join(f"- {r}" for r in XIAOXI_STRICTNESS) + RUBRIC_STYLE_TEXT


# ---- 采访提纲（20 分制） ----
# 来源：小野已批改范例（5 份）提炼，采访部分 12 分 / 其他部分 5 分 / 表述及逻辑 3 分
INTERVIEW_ANCHORS = {
    "采访部分": {
        "low": "采访对象少于3个或角色重叠；问题数量不足（每对象<2个）、闭合式问题过多、问题与对象不对应、无追问、无分点编号",
        "mid": "采访对象基本合理但不够精准或略有重叠；问题基本对应但数量偏少、个别闭合式、追问不足、分点不统一",
        "high": "采访对象≥3个且角色区分清晰、精准对应主题；每对象3-4个问题且数量一致、开放型为主、有追问、分点编号清晰、问题与对象一一对应",
    },
    "其他部分": {
        "low": "缺少采访目的或采访人员及配置等必备板块；背景/主题/方式/时间/地点等关键板块缺失或严重不规范",
        "mid": "板块基本齐全但个别不规范：主题切入点窄、方式名称不写全、时间地点笼统、背景数据不足、人员配置缺失",
        "high": "板块齐全规范：背景丰富含关键数据、主题比标题范围大且用冒号式、方式写全且≥2种、时间具体、地点具体且与对象对应、含采访目的与人员配置",
    },
    "表述及逻辑": {
        "low": "语病较多、口语化严重、表述模糊需二次确认、板块顺序混乱",
        "mid": "少量语病/口语化/简称不写全，不影响理解",
        "high": "书面规范、陈述句为主、表述具体明确、板块顺序合理（背景-目的-主题-方式-时间-地点-对象-问题）",
    },
}
INTERVIEW_STRICTNESS = [
    "严格度：考研阅卷水平，宁严勿松；宁可给低分并说明原因，不可为鼓励而虚高。",
    "硬性校准：采访对象及问题是核心（12分），对象<3个或问题严重不足 → 采访部分 ≤6/12。",
    "缺少采访目的或采访人员及配置 = 其他部分扣分（这是采访提纲必备内容，也是考试重要给分点）。",
    "采访方式名称不写全（如只写'现场'）、时间地点笼统模糊 = 其他部分扣分并红评指出。",
    "闭合式问题过多（需受访者直接回答是/否）→ 采访部分扣分，建议多开放型问题。",
    "法无定法：对象精准、问题充实有追问、模板痕迹不明显，即使不按标准答案也应给高分。",
    "拉开差距：好提纲敢给 16+，有硬伤敢给 ≤10。",
]
INTERVIEW_RUBRIC_TEXT = "\n## 评分锚点（采访部分满分 12，其他部分满分 5，表述及逻辑满分 3）\n" + "\n".join(
    f"- 【{dim}】低：{levels['low']}；中：{levels['mid']}；高：{levels['high']}"
    for dim, levels in INTERVIEW_ANCHORS.items()
) + "\n\n## 严格度校准\n" + "\n".join(f"- {r}" for r in INTERVIEW_STRICTNESS) + RUBRIC_STYLE_TEXT


# ---- 策划（采访提纲/报道提纲，40 分制） ----
# 来源：老师批改范例（02策划 v26290）提炼
# 维度满分：采访主题与目的 8 / 采访对象 10 / 采访问题 15 / 结构完整性 7
PLAN_ANCHORS = {
    "采访主题与目的": {
        "low": "主题不聚焦、与事件脱节；目的缺失或笼统，未分层递进",
        "mid": "主题基本聚焦但表述过长/不够精简；目的有但层次不清或角度单一",
        "high": "主题聚焦具体、比标题范围大；目的分层递进（还原过程→追问责任→揭示漏洞→人文关怀），角度清晰",
    },
    "采访对象": {
        "low": "对象少于3个或角色重叠；对象与主题不对应；遗漏关键主体（如专家、监管方）",
        "mid": "对象基本合理但不够精准；个别对象与主题关联弱；遗漏个别重要角色",
        "high": "对象≥3个且角色区分清晰、精准对应主题；覆盖事件各方（涉事方/监管方/专家/家属/民众）；专家作为独立第三方",
    },
    "采访问题": {
        "low": "问题数量不足（每对象<2个）；闭合式问题过多；问题与对象不对应；无追问；问题偏题",
        "mid": "问题基本对应但数量偏少；个别闭合式；追问不足；个别问题与主题关联弱",
        "high": "每对象3-4个问题且数量一致；开放型为主、有追问；问题与对象一一对应；问题紧扣主题、有深度",
    },
    "结构完整性": {
        "low": "缺少必备板块（背景/主题/目的/对象/问题等）；板块顺序混乱",
        "mid": "板块基本齐全但个别不规范（如背景过长、方式名称不写全、时间地点笼统）",
        "high": "板块齐全规范：背景-主题-目的-准备-采访计划-注意事项顺序合理；背景精简、主题明确、问题为核心",
    },
}
PLAN_STRICTNESS = [
    "严格度：考研阅卷水平，宁严勿松；宁可给低分并说明原因，不可为鼓励而虚高。",
    "硬性校准：采访问题是核心（15分），问题严重不足或大量闭合式 → 采访问题 ≤9/15。",
    "采访对象遗漏专家/监管方等关键角色 = 采访对象扣分（独立第三方能给出客观分析）。",
    "采访背景及主题长篇大论、语言不精简 = 采访主题与目的扣分（考场上写不完的风险）。",
    "采访方式名称不写全（如只写'现场'）、时间地点笼统 = 结构完整性扣分。",
    "法无定法：对象精准、问题充实有追问、模板痕迹不明显，即使不按标准答案也应给高分。",
    "拉开差距：好提纲敢给 32+，有硬伤敢给 ≤24。",
]
PLAN_RUBRIC_TEXT = "\n## 评分锚点（采访主题与目的满分 8，采访对象满分 10，采访问题满分 15，结构完整性满分 7）\n" + "\n".join(
    f"- 【{dim}】低：{levels['low']}；中：{levels['mid']}；高：{levels['high']}"
    for dim, levels in PLAN_ANCHORS.items()
) + "\n\n## 严格度校准\n" + "\n".join(f"- {r}" for r in PLAN_STRICTNESS) + RUBRIC_STYLE_TEXT


# ============================================================
# Prompt 模板
# ============================================================

SYSTEM_PROMPT_TEMPLATE = """你是新闻传播学考研助教「xxxx」，正在批改学生的{homework_type}作业。你严格但温暖，语气亲切自然，像真人老师一样点评。

## 批改优先级（从高到低，先解决最重要的）
1. **论述结构（最高）**：总论点是否鲜明、分论点是否齐全且围绕总论点、每段是否做到"事实→观点收束"闭环、是否像议论文只讲事不讲评、段落衔接是否自然。这类问题用段落级红色评语指出并给改法。
2. **内容事实**：人名/数据/事实是否准确、有无主观臆断。用句子级红色评语+黄色高亮。
3. **语言表达（最后）**：搭配不当、语病、错别字。小问题一句带过，不要喧宾夺主。

## 批改流程
1. 【整体诊断】先看论述结构，判断哪些段落需要段落级评语（结构/逻辑/立意问题）；只在整段确有明显问题时用，不要每段都加。
2. 【逐句批注】对每一句话（包括标题）都给出评语：绿色肯定或红色纠错，不能遗漏。写得好的句子简短肯定并点出好在哪；**问题句不受字数限制，必须展开讲透**（指出问题 + 给具体改法 + 必要时补原因/影响）。
3. 【总评】最后输出评分和鼓励语。

## 评语写作规范（必须遵守）
- 全中文标点：禁止任何半角英文标点（, . ? ! : ; " ' ( ) ~ 等），一律用中文标点（，。？！：；…——（）～）
- 不用引号强调词：应写"应为屠呦呦"，不要写"应为'屠呦呦'"
- 绿色评语简短自然但要点出好在哪（"可以的，数据用得到位。"）；严禁"承接上文/总结分论点/继续叙述/分析论证/收束得当/点明坚持/结构完整/论证有力"这类套话
- 红色评语**不受字数限制**，问题句必须展开：指出问题 + 给具体改法 + 必要时补一句原因/影响
- **禁止空话**：评语必须具体，落到原文的用词、结构、逻辑或材料事实上；严禁"有待加强""继续努力""整体不错"这类无信息量的套话
- 称呼克制：不要每句都以"宝子"开头

## 红绿隔离
- 同一句话只能有一种评语
- 同一段内：仅当该段**确实**加了段落级红色评语（整段结构/逻辑/立意有实质问题）时，才跳过段内句子的绿色肯定，避免红绿并存矛盾
- **不要滥用段落级红批**：只有整段确有实质问题才用；普通段落正常给段内句子绿色肯定

## 输出格式（严格 JSON，不要任何额外文字）
```json
{{
  "total_score": 33,
  "dimensions": {dimensions_json},
  "inline_comments": [
    {{
      "type": "paragraph_red",
      "paragraph_index": 3,
      "anchor_text": "该段唯一短句（可选，用于定位）",
      "comment": "宝子，这段只有事实没有观点收束哦。"
    }},
    {{
      "type": "paragraph_green",
      "paragraph_index": 2,
      "anchor_text": "总论点句",
      "comment": "总论点一针见血，不错的。"
    }},
    {{
      "type": "green",
      "paragraph_index": 3,
      "sentence_text": "原文句子（必须与原文逐字一致）",
      "comment": "标题简明扼要，核心事实把握精准。"
    }},
    {{
      "type": "red",
      "paragraph_index": 4,
      "sentence_text": "有问题的原文片段",
      "comment": "宝子，这里应为屠呦呦哦。"
    }},
    {{
      "type": "strikethrough",
      "paragraph_index": 5,
      "sentence_text": "要加删除线的片段"
    }},
    {{
      "type": "underline",
      "paragraph_index": 3,
      "sentence_text": "写得好的原文句子"
    }}
  ],
  "final_review": {{
    "score_breakdown": {breakdown_example_json},
    "encouragement_paragraphs": ["鼓励段落1", "鼓励段落2", "鼓励段落3"],
    "signature": "批改人：xxxx"
  }}
}}
```

## 批注类型说明
| type | 含义 | 操作 | 关键参数 |
|------|------|------|----------|
| `paragraph_red` | 段落级纠错评语 | 段尾追加红色评语【评语】 | comment, paragraph_index, anchor_text(可选) |
| `paragraph_green` | 段落级肯定评语 | 段尾追加绿色评语【评语】 | comment, paragraph_index, anchor_text(可选) |
| `green` | 句子级肯定 | 原文加绿色下划线+绿色评语【评语】 | sentence_text(原文逐字), comment |
| `red` | 句子级纠错 | 原文加黄色高亮+红色评语【评语】 | sentence_text(原文片段), comment |
| `strikethrough` | 删除标记 | 原文加删除线 | sentence_text(原文片段) |
| `underline` | 肯定标记 | 原文加绿色下划线（不追加评语） | sentence_text(原文逐字) |
| `table_comment` | 表格单元格批注 | 用 Word 批注功能在单元格上添加批语（不改变单元格文字） | table_index, row_index, col_index, comment |

## 技术要点
1. **paragraph_index 是全文段落索引**（即用户 prompt 中 [段落 i] 的 i），与文档实际段落一致
2. **red 类型**：sentence_text 传有问题的原文片段（与原文逐字一致），工具会精确高亮该片段
3. **green 类型**：sentence_text 传完整句子（含标点），与原文逐字一致
4. **strikethrough**：sentence_text 传要加删除线的原文片段
5. **paragraph_red/paragraph_green**：建议同时传 anchor_text（该段的一个唯一短句）
6. **不批改【题目】【材料】等非正文段落**，跳过它们
7. **每句必批**：每一句话（含标题）都要有一条评语，不能遗漏。写得好的给 green/underline，有问题的给 red/strikethrough
8. **总评分数 = 各维度分数之和**
9. **表格批改**：作业含表格时，用 `table_comment` 类型批改需点评的单元格（table_index 从1、row_index/col_index 从0）

## 批改要点
{grading_notes}

## 严格度校准
{strictness_text}
"""


def _build_system_prompt(homework_type: str) -> tuple[str, str, int, list]:
    """根据作业类型构建 system prompt，返回 (system_prompt, grading_notes, score_total, dimensions)"""
    if homework_type == "评论":
        dimensions = COMMENT_DIMENSIONS  # 5维: 立意/标题/论证逻辑/行文结构/语言表达
        score_total = 50
        grading_notes = """## 新闻评论的批改要点
- 总论点是否鲜明、是否在开头点明并与标题呼应（评论的灵魂）
- 分论点是否齐全、是否都围绕总论点展开
- 每段是否做到"事实/事件 → 过渡 → 观点收束"的闭环，还是只讲事不评论
- 是否像高中议论文（只罗列事实、没有"评"）——这是评论大忌
- 引语/开头是否必要；论据是否充分；论证有无跳跃；结尾是否有力
- 语言是否客观中立、符合评论语体
- 法无定法：立意好、论证漂亮，即使不按标准答案也应给分"""
        strictness_text = COMMENT_RUBRIC_TEXT
    elif homework_type == "消息":
        dimensions = ["标题", "导语", "逻辑清晰、事实清楚", "语言通顺、紧扣材料"]
        score_total = 40
        grading_notes = """## 消息改写的批改要点
- 标题是否准确统摄全文、是否为复合标题格式
- 导语是否按倒金字塔结构放最新核心事件
- 数据是否准确、有无遗漏关键数据
- 机构/人名简称是否规范
- 语言是否客观中立，有无主观色彩
- 结构是否合理、有无倒金字塔意识"""
        strictness_text = XIAOXI_RUBRIC_TEXT
    elif homework_type == "采访":
        dimensions = INTERVIEW_DIMENSIONS  # 3维: 采访部分/其他部分/表述及逻辑
        score_total = INTERVIEW_SCORE_TOTAL  # 20 分制
        grading_notes = """## 采访提纲的批改要点
- **采访对象及问题是核心（12分）**：对象≥3个且角色区分清晰、不重叠、精准对应主题；每对象3-4个问题且数量一致、开放型为主（避免闭合式）、有追问、分点编号清晰、问题与对象一一对应
- **采访背景**：是否丰富、是否包含关键数据（如开始年份、车站数量等）
- **采访主题**：是否比标题范围大、是否用冒号式（如"聚焦夏日民生：成都地铁五年车站纳凉便民服务全方位探索"）
- **采访目的**：必不可少，建议单独列出或并入背景
- **采访方式**：名称要写全（"现场采访"而非"现场"）、至少两种
- **采访时间**：要具体明确（具体日期、时间段）
- **采访地点**：要具体（具体站点、纳凉区、车厢等）、与采访对象对应、3-5个
- **采访人员及配置**：必不可少，可与器材设备合并为一个大点
- **采访步骤/可能碰到的问题/解决办法**：20分提纲可省略，写了要写全
- 法无定法：对象精准、问题充实有追问、模板痕迹不明显，即使不按标准答案也应给分"""
        strictness_text = INTERVIEW_RUBRIC_TEXT
    elif homework_type == "策划":
        dimensions = ["采访主题与目的", "采访对象", "采访问题", "结构完整性"]
        score_total = 40
        grading_notes = """## 策划（采访提纲/报道提纲）的批改要点
- **采访问题是核心（15分）**：每对象3-4个问题且数量一致、开放型为主（避免闭合式）、有追问、问题与对象一一对应、紧扣主题
- **采访对象（10分）**：≥3个且角色区分清晰、精准对应主题；覆盖事件各方（涉事方/监管方/专家/家属/民众）；专家作为独立第三方能给出客观分析
- **采访主题与目的（8分）**：主题聚焦具体、比标题范围大；目的分层递进（还原过程→追问责任→揭示漏洞→人文关怀）
- **结构完整性（7分）**：背景-主题-目的-准备-采访计划-注意事项顺序合理；背景精简、主题明确、问题为核心
- **背景及主题不要长篇大论**：语言要精简，这两部分写太多有考场上写不完的风险
- 法无定法：对象精准、问题充实有追问、模板痕迹不明显，即使不按标准答案也应给分"""
        strictness_text = PLAN_RUBRIC_TEXT
    else:
        dimensions = ["标题", "导语", "逻辑清晰、事实清楚", "语言通顺、紧扣材料"]
        score_total = 40
        grading_notes = "请按消息改写标准批改。"
        strictness_text = XIAOXI_RUBRIC_TEXT

    # 各维度满分：采访作业为 12/5/3，策划为 8/10/15/7，其余类型统一 10 分
    if homework_type == "采访":
        dim_max = {"采访部分": 12, "其他部分": 5, "表述及逻辑": 3}
    elif homework_type == "策划":
        dim_max = {"采访主题与目的": 8, "采访对象": 10, "采访问题": 15, "结构完整性": 7}
    else:
        dim_max = {d: 10 for d in dimensions}

    dims_json = json.dumps(
        {d: {"score": 0, "max": dim_max.get(d, 10), "comment": ""} for d in dimensions},
        ensure_ascii=False, indent=2,
    )

    # score_breakdown 示例（按作业类型维度生成）
    breakdown_items = [f"{d}：X/{dim_max.get(d, 10)}" for d in dimensions]
    breakdown_items.append(f"总分：X/{score_total}")
    breakdown_example_json = json.dumps(breakdown_items, ensure_ascii=False)

    system_prompt = SYSTEM_PROMPT_TEMPLATE.format(
        homework_type=homework_type,
        dimensions_json=dims_json,
        score_total=score_total,
        grading_notes=grading_notes,
        strictness_text=strictness_text,
        breakdown_example_json=breakdown_example_json,
    )
    return system_prompt, grading_notes, score_total, dimensions


# ============================================================
# 诊断轮（诊断先行：先全局诊断，再带诊断逐句批注）
# ============================================================
DIAGNOSE_SYSTEM_PROMPT = """你是新闻传播学考研助教「xxxx」，批改学生的{homework_type}作业。
在正式批注前，你需要先对全文做一次**全局诊断**，为后续逐句批注提供方向。

## 你的任务
通读学生作业全文 + 参考答案，逐段诊断，产出结构化诊断 JSON。
诊断要**深刻、具体**：指出每段的问题类型、给出具体改法、点名材料里可引用的人名/数据。

## 诊断维度（按作业类型）
{grading_notes}

## 诊断要求
1. **overall**：立意判断（是否鲜明/是否偏题/有无"评"）、全文最核心的 1-2 个问题、给总评用的方向素材
2. **paragraphs**：逐段诊断（答案正文的每一段都要覆盖），每段一个条目：
   - index: 段落全文索引
   - 优点: 这段写得好的地方（每段都要写，哪怕只有一点，如"数据用得到位""转折自然"）
   - 问题类型: 如 转述材料/缺观点收束/跑题/事实错误/语言问题/结构问题...；**若该段整体良好、没有实质问题，填"整体良好"**
   - 具体改法: 有问题才写，给学生可操作的具体修改建议（要具体到点名材料里的人名/数据，如"材料里是图灵奖得主姚期智，建议写明"）；整体良好则留空
   - 评语方向: 这段的批语想让学生改什么（一句话）；整体良好则填肯定方向
3. **平衡诊断（重要）**：不要为了"每段都有条目"而硬找问题。只有**确实存在**的实质问题才标记；小瑕疵（个别用词、标点）不要升级为段落问题。**好段就如实写"整体良好"**，不要吹毛求疵。
4. **红绿隔离**：仅当某段**确实需要段落级红色评语**（整段的结构/逻辑/立意有实质问题）时，才在诊断中标记该段为"需段落级红批"；普通段落不要标记，段内句子仍可正常给绿色肯定。

## 输出格式（严格 JSON，不要任何额外文字）
```json
{{
  "overall": {{
    "立意判断": "...",
    "核心问题": "全文最大的1-2个问题",
    "总评方向": "给总评用的素材（整体进步方向/最值得肯定处/最需改进的一个方向）"
  }},
  "paragraphs": [
    {{
      "index": 3,
      "优点": "数据用得到位，把发展机遇说实了",
      "问题类型": "转述材料/缺观点收束",
      "具体改法": "建议把重心落在你的判断上，如：这四项主张切中了治理的哪些命门",
      "评语方向": "让学生从复述材料转为给出自己的判断"
    }},
    {{
      "index": 5,
      "优点": "转折自然，观点明确",
      "问题类型": "整体良好",
      "具体改法": "",
      "评语方向": "肯定转折与观点表达"
    }}
  ]
}}
```
"""


def _build_diagnose_prompt(
    homework_text: str,
    paragraphs: list[str],
    pdf_body: str,
    answer_indexes: Optional[list[int]] = None,
    grading_notes: str = "",
    tables_data: Optional[list[dict]] = None,
) -> str:
    """
    构建诊断轮 user prompt：答案正文 + 参考答案 + 表格内容（参考答案只在诊断轮给药，批注轮同窗口自动可见）。
    """
    answer_indexes = answer_indexes or list(range(len(paragraphs)))

    para_lines = []
    for i, p in enumerate(paragraphs):
        if not p.strip() or i not in answer_indexes:
            continue
        para_lines.append(f"[段落 {i}] {p}")

    homework_with_index = "\n\n".join(para_lines)

    # 表格内容（如有）单独列出，供 LLM 诊断
    tables_text = ""
    if tables_data:
        table_lines = []
        for t in tables_data:
            table_lines.append(f"[表格 {t['index']}]")
            for ri, row in enumerate(t["rows"]):
                row_text = " | ".join(row)
                table_lines.append(f"  行{ri}: {row_text}")
        tables_text = "\n\n## 学生作业表格（表格内容也是答案正文的一部分，需诊断）\n" + "\n".join(table_lines)

    return (
        f"## 参考答案\n{pdf_body}\n\n"
        f"## 学生作业（带全文段落索引，只含答案正文，skip题目/材料/说明）\n{homework_with_index}\n\n"
        f"{tables_text}\n\n"
        "请先做全局诊断，输出结构化诊断 JSON。不要输出批注，诊断即可。"
    )


def _build_comment_prompt(
    homework_text: str,
    paragraphs: list[str],
    answer_indexes: Optional[list[int]] = None,
    tables_data: Optional[list[dict]] = None,
) -> str:
    """
    构建批注轮 user prompt：答案正文 + 引导其参考上文诊断。
    同窗口下，参考答案与诊断已在历史消息中，无需重复给药。
    """
    answer_indexes = answer_indexes or list(range(len(paragraphs)))

    para_lines = []
    sentence_count = 0
    for i, p in enumerate(paragraphs):
        if not p.strip() or i not in answer_indexes:
            continue
        para_lines.append(f"[段落 {i}] {p}")
        sents = [s for s in re.split(r'(?<=[。；！？])', p) if s.strip()]
        sentence_count += len(sents)

    homework_with_index = "\n\n".join(para_lines)

    # 表格内容（如有）单独列出，供 LLM 批注表格单元格
    tables_text = ""
    table_requirement = ""
    if tables_data:
        table_lines = []
        for t in tables_data:
            table_lines.append(f"[表格 {t['index']}]")
            for ri, row in enumerate(t["rows"]):
                row_text = " | ".join(row)
                table_lines.append(f"  行{ri}: {row_text}")
        tables_text = "\n\n## 学生作业表格（表格内容也是答案正文的一部分，必须批改）\n" + "\n".join(table_lines)
        table_requirement = (
            "## 表格批改要求\n"
            "学生作业包含上述表格，表格内容也是答案正文的一部分，必须批改。"
            "对表格中需要点评的单元格，使用 type 为 `table_comment` 的批注，"
            "参数：`table_index`（表格序号，从1开始）、`row_index`（行号，从0开始）、"
            "`col_index`（列号，从0开始）、`comment`（评语内容）。\n\n"
        )

    return (
        f"## 学生作业（带全文段落索引，只含答案正文）\n{homework_with_index}\n\n"
        f"{tables_text}\n\n"
        f"## 批改要求\n"
        f"学生作业共 {sentence_count} 句话，你的 inline_comments 必须包含至少 {sentence_count} 条评语，"
        f"每句话至少一条，不能遗漏任何一句（包括标题）。\n\n"
        f"{table_requirement}"
        "## 索引说明\n"
        "所有 paragraph_index 使用**全文段落索引**（即上面 [段落 i] 中的 i），不是答案相对索引。\n\n"
        "## red 类型 sentence_text 粒度（必须遵守）\n"
        "red 类型的 sentence_text 只传**有问题的精确片段**（如错词\"是老化\"、\"踉跄跄\"），"
        "**不要传整句**。工具会精确高亮该片段，评语紧跟其后。若整句都有问题，才传整句。\n\n"
        "## 平衡批改（必须遵守）\n"
        "1. **逐句平衡**：好句给绿色肯定（简短自然，点出好在哪），问题句给红色纠错。不要只盯着问题找茬。\n"
        "2. **每段至少一条绿色**：除非某段确实一无是处，否则每段都要有至少一条绿色肯定（标题、总论点、转折、收束、数据运用等都可肯定）。\n"
        "3. **参考诊断全部内容**：上一步诊断的 overall（立意判断/核心问题/总评方向）与逐段诊断（优点/问题类型/具体改法/评语方向）**全部都要参考**，把诊断里的具体改法、点名的材料事实落到对应句子的批语上；诊断未提的句子仍按实际质量批改；诊断标记为'整体良好'的段落，正常给绿色肯定。\n"
        "4. **红绿隔离**：仅当某段确实加了段落级红色评语（整段结构/逻辑/立意有实质问题）时，才跳过该段内句子的绿色肯定；普通段落不要跳过。\n"
        "5. **小瑕疵不升级**：个别用词、标点等小问题一句带过即可，不要整段标红。标点问题不要求区分中英文以及全半角。\n"
        "6. **问题句展开**：红色评语不受字数限制，必须展开讲透——指出问题 + 给具体改法 + 必要时补原因/影响。\n"
        "7. **禁止空话**：评语必须具体，落到原文的用词、结构、逻辑或材料事实上；严禁'有待加强''继续努力''整体不错'这类无信息量的套话。\n\n"
        "请结合上一步诊断的全部内容，逐句平衡批注，输出批注 JSON + 总评，不要任何额外文字。"
    )


# ============================================================
# LLM 调用
# ============================================================

def _get_client() -> OpenAI:
    """获取 DeepSeek 客户端"""
    api_key = os.getenv("ARK_API_KEY") or ARK_API_KEY
    if not api_key:
        raise ValueError(
            "❌ 未找到 API Key！请在项目根目录创建 .env 文件（内容：ARK_API_KEY=你的密钥），"
            "或设置系统环境变量 ARK_API_KEY"
        )
    return OpenAI(base_url=DEFAULT_BASE_URL, api_key=api_key)


def _try_fix_truncated_json(text: str) -> str:
    """尝试修复截断的 JSON：补全未闭合的字符串与括号，并丢弃末尾残留"""
    if not text.strip():
        return text

    first_brace = text.find('{')
    if first_brace == -1:
        return text
    text = text[first_brace:]

    # 去掉 markdown 围栏/末尾杂散：只保留从 { 到最后一个 } 的部分（丢弃不完整尾部）
    last_brace = text.rfind('}')
    if last_brace > 0:
        text = text[:last_brace + 1]

    # 结构性修复：扫描并补全未闭合的字符串与括号
    out = []
    in_string = False
    escape = False
    stack = []
    for ch in text:
        if in_string:
            out.append(ch)
            if escape:
                escape = False
            elif ch == '\\':
                escape = True
            elif ch == '"':
                in_string = False
            continue
        if ch == '"':
            in_string = True
            out.append(ch)
        elif ch in '{[':
            stack.append(ch)
            out.append(ch)
        elif ch in '}]':
            if stack:
                stack.pop()
            out.append(ch)
        else:
            out.append(ch)
    if in_string:
        out.append('"')  # 补上未闭合字符串的引号
    while stack:
        opener = stack.pop()
        out.append('}' if opener == '{' else ']')
    return ''.join(out)


# ============================================================
# 清洗阶段：用 LLM 从原始文档提取「答案正文」段落（带全文索引）
# ============================================================
CLEAN_SYSTEM_PROMPT = """你是文档清洗助手。给定一份学生作业的全文段落（每段带全文索引 [i]），
请识别出「学生答案正文」的段落，排除题目、材料、说明、模板等非正文内容。

## 判定规则
- 【题目】【材料】及其正文、模板说明行、评分说明等 → 排除
- 学生实际作答的正文（评论/消息/答案）→ 保留，包括标题
- 空段落 → 排除

## 输出格式（严格 JSON，不要任何额外文字）
```json
{
  "answer_paragraph_indexes": [3, 5, 6, 7, 8]
}
```
answer_paragraph_indexes 是保留段落的**全文索引**（与输入 [i] 一致），按从小到大排列。
"""


def _extract_answer_with_llm(docx_path: str) -> tuple[list[int], list[str]]:
    """
    用 LLM 从原始文档提取答案正文段落的全文索引。
    相比启发式 extract_answer_part，对任意格式的学生文档更鲁棒。
    返回 (answer_indexes, full_paragraphs)：
      - answer_indexes: 答案正文段落的全文索引（升序）
      - full_paragraphs: 全文段落文本列表（与索引对齐）
    """
    doc = Document(docx_path)
    paras = [p.text for p in doc.paragraphs]

    # 构造带索引的段落列表（含空段标记，便于模型判断边界）
    lines = []
    for i, t in enumerate(paras):
        label = t.strip() if t.strip() else "<空段>"
        lines.append(f"[{i}] {label}")
    full_text = "\n".join(lines)

    user_prompt = (
        f"## 学生作业全文段落（带全文索引）\n{full_text}\n\n"
        "请识别学生答案正文的段落索引，排除题目/材料/说明/模板/空段。"
    )

    result = _call_llm(CLEAN_SYSTEM_PROMPT, user_prompt, max_retries=3)
    idxs = result.get("answer_paragraph_indexes", [])
    # 过滤：只保留有效索引，且按升序
    valid = sorted(i for i in idxs if isinstance(i, int) and 0 <= i < len(paras))
    return valid, paras


def _extract_tables(docx_path: str) -> list[dict]:
    """
    从 docx 中提取所有表格内容，供 LLM 批改表格单元格。
    返回 [{"index": 1, "rows": [[...], ...]}, ...]，表格序号从 1 开始。
    """
    doc = Document(docx_path)
    tables_data = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables_data.append({"index": ti + 1, "rows": rows})
    return tables_data


def _call_llm(
    system_prompt: str,
    user_prompt: str,
    max_retries: int = 5,
    messages: Optional[list] = None,
) -> dict:
    """
    调用 LLM 并解析 JSON 结果。

    Args:
        system_prompt: system 消息
        user_prompt: user 消息
        max_retries: 最大重试次数
        messages: 可选，多轮对话历史（同窗口复用）。
                  若提供，则以 messages 为基础，在其后追加本次 user 消息。
                  否则新建 [system, user] 两轮。

    Returns:
        解析后的 JSON dict。
    """
    client = _get_client()

    for attempt in range(max_retries):
        try:
            # 构造消息列表：复用历史（同窗口）或新建
            if messages:
                full_messages = list(messages) + [
                    {"role": "user", "content": user_prompt},
                ]
            else:
                full_messages = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ]

            _api_kwargs = dict(
                model=DEFAULT_MODEL,
                messages=full_messages,
                temperature=TEMPERATURE,
                max_tokens=MAX_TOKENS,
                response_format={"type": "json_object"},
            )
            if REASONING_EFFORT:
                _api_kwargs["extra_body"] = {"reasoning_effort": REASONING_EFFORT}

            response = client.chat.completions.create(**_api_kwargs)
            result_text = response.choices[0].message.content

            # 解析 JSON（先尝试完整解析，失败则尝试修复截断）
            try:
                result = json.loads(result_text)
            except json.JSONDecodeError as e:
                # 修复截断：去掉末尾不完整部分
                fixed = _try_fix_truncated_json(result_text)
                if fixed != result_text:
                    print(f"   ⚠️  JSON 被截断，已尝试修复（{len(result_text)}→{len(fixed)} 字符）")
                    try:
                        result = json.loads(fixed)
                    except json.JSONDecodeError:
                        # 再试一次 markdown 包裹
                        json_match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", fixed, re.DOTALL)
                        if json_match:
                            result = json.loads(json_match.group(1))
                        else:
                            raise e
                else:
                    json_match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", result_text, re.DOTALL)
                    if json_match:
                        result = json.loads(json_match.group(1))
                    else:
                        raise

            # 验证必要字段
            if "inline_comments" not in result:
                result["inline_comments"] = []
            if "final_review" not in result:
                result["final_review"] = {}
            if "dimensions" not in result:
                result["dimensions"] = {}
            if "total_score" not in result:
                result["total_score"] = 0

            return result

        except Exception as e:
            if attempt < max_retries - 1:
                wait = 5 * (attempt + 1)
                print(f"   ⚠️  LLM 调用失败（第{attempt+1}次）: {e}")
                print(f"   ⏳ {wait}秒后重试...")
                time.sleep(wait)
            else:
                raise RuntimeError(f"LLM 调用失败（已重试{max_retries}次）: {e}")


# ============================================================
# 批注注入（Phase 2）
# ============================================================

def _inject_comments(current_docx: str, comments: list) -> str:
    """逐条注入评语到文档"""
    if not comments:
        print("   📭 没有需要批注的评论")
        return current_docx

    ok = 0
    failed = 0
    skipped = 0

    # 结构优先级/红绿隔离：有段落级红色评语（结构问题）的段落，跳过其内句子级绿色肯定
    red_para_indexes = {
        c.get("paragraph_index") for c in comments if c.get("type") == "paragraph_red"
    }

    for i, c in enumerate(comments):
        ctype = c.get("type", "")
        para_idx = c.get("paragraph_index", 0)
        comment_text = c.get("comment", "")
        sentence_text = c.get("sentence_text", "")
        anchor_text = c.get("anchor_text", "")

        # 红绿隔离：段落级红色评语所在段落，跳过其内句子级绿色肯定（green/underline）
        if ctype in ("green", "underline") and para_idx in red_para_indexes:
            print(f"   ⏭️  [{i}] 段落{para_idx}已有段落级红色评语（结构问题），跳过段内绿色肯定")
            skipped += 1
            continue

        try:
            if ctype == "paragraph_red":
                result = tool_add_paragraph_comment(
                    docx_path=current_docx,
                    comment=comment_text,
                    paragraph_index=para_idx,
                    color="red",
                    anchor_text=anchor_text or None,
                )
            elif ctype == "paragraph_green":
                result = tool_add_paragraph_comment(
                    docx_path=current_docx,
                    comment=comment_text,
                    paragraph_index=para_idx,
                    color="green",
                    anchor_text=anchor_text or None,
                )
            elif ctype == "green":
                if not sentence_text:
                    print(f"   ⚠️  [{i}] green 类型缺少 sentence_text，跳过")
                    failed += 1
                    continue
                result = tool_add_green_sentence(
                    docx_path=current_docx,
                    sentence_text=sentence_text,
                    comment=comment_text,
                    paragraph_index=para_idx,
                )
            elif ctype == "red":
                if not sentence_text:
                    print(f"   ⚠️  [{i}] red 类型缺少 sentence_text，跳过")
                    failed += 1
                    continue
                result = tool_add_red_sentence(
                    docx_path=current_docx,
                    sentence_text=sentence_text,
                    comment=comment_text,
                    paragraph_index=para_idx,
                )
            elif ctype == "strikethrough":
                if not sentence_text:
                    print(f"   ⚠️  [{i}] strikethrough 类型缺少 sentence_text，跳过")
                    failed += 1
                    continue
                result = tool_add_strikethrough(
                    docx_path=current_docx,
                    sentence_text=sentence_text,
                    paragraph_index=para_idx,
                )
            elif ctype == "underline":
                # 下划线用绿色句子的标记部分，但不追加评语
                if not sentence_text:
                    print(f"   ⚠️  [{i}] underline 类型缺少 sentence_text，跳过")
                    failed += 1
                    continue
                result = tool_add_green_sentence(
                    docx_path=current_docx,
                    sentence_text=sentence_text,
                    comment="",  # 空评语 → 只加下划线不追加评语
                    paragraph_index=para_idx,
                )
            elif ctype == "table_comment":
                # 表格单元格批注（Word 批注功能）
                if not comment_text:
                    print(f"   ⚠️  [{i}] table_comment 类型缺少 comment，跳过")
                    failed += 1
                    continue
                result = tool_add_table_comment(
                    docx_path=current_docx,
                    comment=comment_text,
                    table_index=c.get("table_index", 1),
                    row_index=c.get("row_index", 0),
                    col_index=c.get("col_index", 0),
                )
            else:
                print(f"   ⚠️  [{i}] 未知类型: {ctype}，跳过")
                failed += 1
                continue

            # 更新 current_docx
            if "docx_path" in result:
                current_docx = result["docx_path"]

            if result.get("error"):
                print(f"   ❌ [{i}] {ctype}: {result['error'][:60]}")
                failed += 1
            elif result.get("skipped"):
                label = {"green": "🟢", "red": "🔴", "paragraph_red": "🔴段", "paragraph_green": "🟢段",
                         "strikethrough": "删除", "underline": "下划线", "table_comment": "📋表格"}.get(ctype, ctype)
                print(f"   ⏭️  [{i}] {label} 段落{para_idx}: {result.get('reason', '跳过')[:40]}")
                skipped += 1
            else:
                label = {"green": "🟢", "red": "🔴", "paragraph_red": "🔴段", "paragraph_green": "🟢段",
                         "strikethrough": "删除线", "underline": "下划线", "table_comment": "📋表格"}.get(ctype, ctype)
                print(f"   ✅ [{i}] {label} 段落{para_idx}")
                ok += 1

        except Exception as e:
            print(f"   ❌ [{i}] {ctype} 异常: {e}")
            failed += 1

    print(f"\n   📊 批注注入结果: ✅ {ok} 成功 / ⏭️  {skipped} 跳过 / ❌ {failed} 失败 / 共 {len(comments)} 条")
    return current_docx


# ============================================================
# 主入口
# ============================================================

def run_single_shot_grader(
    answer_pdf_path: str,
    homework_docx_path: str,
    homework_type: str = "通用",
    output_docx_path: Optional[str] = None,
    author: str = "xxxx",
) -> dict:
    """
    运行 Single-Shot 批改流程

    Args:
        answer_pdf_path: 参考答案 PDF 路径
        homework_docx_path: 学生作业 docx 路径
        homework_type: 作业类型（"评论" / "消息" / "通用"）
        output_docx_path: 输出路径（可选，默认自动生成）
        author: 批改人署名

    Returns:
        dict: 批改结果汇总
    """
    SEP = "=" * 60
    print(SEP)
    print("Word Agent 智能批改系统 (Single-Shot 模式)")
    print(SEP)

    # ---- 第1步：读取数据 ----
    print("\n📖 第1步：读取参考答案和学生作业...")

    # 读取参考答案 PDF
    pdf_data = tool_read_answer_pdf(answer_pdf_path)
    pdf_lines = pdf_data["text"].splitlines()
    pdf_body = "\n".join(
        line for line in pdf_lines if line.strip() and "作业" not in line[:4]
    )[:4000]
    print(f"   ✅ 参考答案: {pdf_data.get('lines', 0)} 行")

    # 读取学生作业（全文段落，用于清洗与批改）
    hw_data = tool_read_homework(homework_docx_path)
    print(f"   ✅ 学生作业: {hw_data['paragraph_count']} 段, {hw_data['sentence_count']} 句")

    # ---- 第2步：标准化文档 ----
    print("\n🔧 第2步：标准化文档 run 结构...")
    current_docx = normalize_document_runs(homework_docx_path)
    print(f"   ✅ 已标准化")

    # ---- 第2.5步：LLM 清洗，提取答案正文段落（全文索引） ----
    print("\n🧹 第2.5步：LLM 清洗，识别答案正文...")
    answer_indexes, full_paragraphs = _extract_answer_with_llm(homework_docx_path)
    if not answer_indexes:
        print("   ⚠️  清洗未识别到答案正文，回退为全文批改")
        answer_indexes = list(range(len(full_paragraphs)))
    print(f"   ✅ 答案正文段落: {answer_indexes}")

    # 第2.6步：抽取表格内容（采访提纲等含表格的作业）
    tables_data = _extract_tables(homework_docx_path)
    if tables_data:
        print(f"   ✅ 发现 {len(tables_data)} 个表格，将一并批改")

    # ---- 第3步：调用 LLM（诊断先行 + 同窗口批注） ----
    print("\n🤖 第3步：调用 DeepSeek 批改（诊断先行，同窗口两轮）...")

    system_prompt, grading_notes, score_total, dimensions = _build_system_prompt(homework_type)

    # 3a. 诊断轮：全局诊断，产出结构化 JSON（参考答案只在此给药）
    print("   第3a步：全局诊断中...")
    diagnose_system = DIAGNOSE_SYSTEM_PROMPT.format(
        homework_type=homework_type,
        grading_notes=grading_notes,
    )
    diagnose_prompt = _build_diagnose_prompt(
        hw_data["text"],
        full_paragraphs,
        pdf_body,
        answer_indexes=answer_indexes,
        grading_notes=grading_notes,
        tables_data=tables_data,
    )
    diagnose_result = _call_llm(diagnose_system, diagnose_prompt, max_retries=3)
    paragraphs_diag = diagnose_result.get("paragraphs", [])
    overall_diag = diagnose_result.get("overall", {})
    print(f"   ✅ 诊断完成: {len(paragraphs_diag)} 段诊断")

    # 3b. 批注轮：同一窗口，带诊断历史逐句批注 + 总评
    print("   第3b步：基于诊断逐句批注中...")
    # 构造同窗口历史消息：system + 诊断轮(user+assistant)
    history = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": diagnose_prompt},
        {"role": "assistant", "content": json.dumps(diagnose_result, ensure_ascii=False)},
    ]
    comment_prompt = _build_comment_prompt(
        hw_data["text"],
        full_paragraphs,
        answer_indexes=answer_indexes,
        tables_data=tables_data,
    )
    result = _call_llm(system_prompt, comment_prompt, messages=history)

    # 显示评分
    total = result.get("total_score", 0)
    dims = result.get("dimensions", {})
    print("\n   📊 批改结果:")
    print("   " + "-" * 30)
    for name, data in dims.items():
        s = data.get("score", 0) if isinstance(data, dict) else 0
        m = data.get("max", 10) if isinstance(data, dict) else 10
        bar = "█" * s + "░" * (m - s)
        print(f"   {name}: {s}/{m}  {bar}")
    print("   " + "-" * 30)
    print(f"   总分: {total}")

    # ---- 第4步：注入评语 ----
    print("\n✏️  第4步：注入批注到文档...")
    comments = result.get("inline_comments", [])
    print(f"   共 {len(comments)} 条批注")
    current_docx = _inject_comments(current_docx, comments)

    # ---- 第5步：追加总评 ----
    print("\n📝 第5步：追加总评...")
    final_review = result.get("final_review", {})
    if final_review and final_review.get("score_breakdown"):
        fr_result = tool_append_final_review(
            docx_path=current_docx,
            score_breakdown=final_review.get("score_breakdown", []),
            disclaimer=final_review.get("disclaimer", ""),
            encouragement_paragraphs=final_review.get("encouragement_paragraphs", []),
            signature=final_review.get("signature", "") or f"批改人：{author}",
        )
        if "docx_path" in fr_result:
            current_docx = fr_result["docx_path"]
        print("   ✅ 总评已追加")
    else:
        print("   ⚠️  未生成总评（JSON 中缺少 final_review.score_breakdown）")

    # ---- 第6步：输出 ----
    print("\n💾 第6步：保存文件...")
    if output_docx_path is None:
        src = Path(homework_docx_path)
        # 当日日期子文件夹（如 2026-8-19），避免文件堆积在 complete_homework 顶层
        now = time.localtime()
        date_dir = f"{now.tm_year}-{now.tm_mon}-{now.tm_mday}"
        output_docx_path = str(src.parent.parent / "complete_homework" / date_dir / f"【已批改】{src.stem}{src.suffix}")

    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)

    # 从 temp 复制到最终输出
    try:
        shutil.copy2(current_docx, output_docx_path)
        print(f"   ✅ 输出文件: {output_docx_path}")
    except Exception as e:
        print(f"   ❌ 复制失败: {e}")
        print(f"   临时文件保留在: {current_docx}")

    # 清理临时文件
    cleanup_temp_files(keep_paths=[current_docx])

    print("\n" + SEP)
    print("✅ 批改完成！")
    print(SEP)

    return {
        "total_score": total,
        "dimensions": dims,
        "output_path": output_docx_path,
        "comments_count": len(comments),
    }