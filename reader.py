"""
Word 文档读取模块
支持读取 .docx 文件的文本内容、表格、段落等
"""

from pathlib import Path
from typing import Optional
from docx import Document


def read_docx(file_path: str) -> str:
    """
    读取 Word 文档的全部文本内容

    Args:
        file_path: Word 文档路径

    Returns:
        文档的文本内容字符串
    """
    doc = _load_document(file_path)
    parts = []

    # 读取段落
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    # 读取表格
    for i, table in enumerate(doc.tables):
        parts.append(f"\n--- 表格 {i + 1} ---")
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            parts.append(row_text)

    return "\n".join(parts)


def read_docx_with_style(file_path: str) -> list[dict]:
    """
    读取 Word 文档，返回带样式的结构化内容

    Args:
        file_path: Word 文档路径

    Returns:
        包含段落样式信息的列表，每项为 {"text": str, "style": str, "type": str}
    """
    doc = _load_document(file_path)
    content = []

    for para in doc.paragraphs:
        if para.text.strip():
            content.append({
                "text": para.text.strip(),
                "style": para.style.name if para.style else "Normal",
                "type": "paragraph"
            })

    for i, table in enumerate(doc.tables):
        content.append({
            "text": f"[表格 {i + 1}]",
            "style": "Table",
            "type": "table_header"
        })
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            content.append({
                "text": row_text,
                "style": "Table Row",
                "type": "table_row"
            })

    return content


def get_paragraphs(file_path: str) -> list[str]:
    """
    获取文档中所有段落文本

    Args:
        file_path: Word 文档路径

    Returns:
        段落文本列表
    """
    doc = _load_document(file_path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def get_tables(file_path: str) -> list[list[list[str]]]:
    """
    获取文档中所有表格数据

    Args:
        file_path: Word 文档路径

    Returns:
        三维列表: [表格][行][列]
    """
    doc = _load_document(file_path)
    tables_data = []

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables_data.append(table_data)

    return tables_data


def extract_answer_part(file_path: str) -> tuple[str, list[str], int]:
    """
    从学生作业中提取【答案】部分，跳过题目和材料

    许多作业文档包含：题目 → 材料 → 答案 三部分。
    此函数只提取「答案」部分的内容用于批改和添加批注，
    同时保持段落索引与 Word 文档实际段落一致以便加批注。

    Args:
        file_path: Word 文档路径

    Returns:
        (answer_text, answer_paragraphs, answer_start_index)
        - answer_text: 答案部分的纯文本
        - answer_paragraphs: 答案部分的段落列表（保持原始段落索引）
        - answer_start_index: 答案部分在全文段落中的起始索引
    """
    doc = _load_document(file_path)
    all_paras = [(i, p.text) for i, p in enumerate(doc.paragraphs)]

    # 寻找【答案】标记
    answer_start = -1
    for i, (_, text) in enumerate(all_paras):
        if "【答案】" in text:
            answer_start = i
            break

    if answer_start == -1:
        # 没有【答案】标记：找【题目】和【材料】的位置
        topic_idx = -1
        material_idx = -1
        for i, (_, text) in enumerate(all_paras):
            if "【题目】" in text:
                topic_idx = i
            if "【材料】" in text:
                material_idx = i

        if material_idx >= 0 or topic_idx >= 0:
            # 有【题目】/【材料】：从它们结束后找第一个空段之后的下一个非空段
            search_start = max(material_idx, topic_idx) + 1
            found_empty = False
            content_start = -1
            for i in range(search_start, len(all_paras)):
                _, text = all_paras[i]
                if not text.strip():
                    found_empty = True
                elif found_empty:
                    content_start = i
                    break
        else:
            # 无【题目】无【材料】无【答案】：整个文档就是学生作答，从第一个非空段开始
            content_start = -1
            for i, (_, text) in enumerate(all_paras):
                if text.strip():
                    content_start = i
                    break

        # 如果没找到空段分隔，回退到 search_start 后的第一个非空段
        if content_start == -1:
            search_start = max(material_idx, topic_idx) + 1 if max(material_idx, topic_idx) >= 0 else 0
            for i in range(search_start, len(all_paras)):
                _, text = all_paras[i]
                if text.strip():
                    content_start = i
                    break
        if content_start > 0:
            answer_paras = [all_paras[i][1] for i in range(content_start, len(all_paras))]
            answer_text = "\n".join(t for t in answer_paras if t.strip())
            return answer_text, answer_paras, content_start
        # 兜底：返回全部
        all_text = "\n".join(t for _, t in all_paras if t.strip())
        return all_text, [t for _, t in all_paras], 0

    # 答案部分所有段落（包括空段落，保持索引对齐）
    answer_paras = [all_paras[i][1] for i in range(answer_start + 1, len(all_paras))]
    answer_text = "\n".join(t for t in answer_paras if t.strip())
    return answer_text, answer_paras, answer_start + 1


def _load_document(file_path: str) -> Document:
    """
    加载 Word 文档，支持路径验证

    Args:
        file_path: Word 文档路径

    Returns:
        Document 对象

    Raises:
        FileNotFoundError: 文件不存在
        ValueError: 文件格式不支持
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    if path.suffix.lower() not in (".docx",):
        raise ValueError(f"不支持的文件格式: {path.suffix}，仅支持 .docx")

    return Document(str(path))