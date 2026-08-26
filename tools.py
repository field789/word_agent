"""
Word Agent 工具集 — 句子级批注
每个工具格式硬编码，模型只需选择工具和参数。
"""

import os
import gc
import shutil
import tempfile
import time
import re
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

from word_agent.reader import extract_answer_part

# ============================================================
# 格式常量 & 工具路径
# ============================================================
COLOR_GREEN = RGBColor(0x00, 0xB0, 0x50)
COLOR_RED   = RGBColor(0xEE, 0x00, 0x00)
COLOR_BLUE  = RGBColor(0x00, 0x70, 0xC0)

_temp_files = []

# 临时目录：D:\Agent\word_agent\temp
TOOLS_DIR = Path(__file__).resolve().parent
TEMP_DIR = TOOLS_DIR / "temp"


def _ensure_temp_dir():
    """确保 temp 目录存在"""
    TEMP_DIR.mkdir(parents=True, exist_ok=True)


def _copy_to_temp(docx_path):
    """复制文档到 word_agent/temp/ 下的唯一子目录"""
    _ensure_temp_dir()
    work_dir = TEMP_DIR / f"work_{int(time.time() * 1000)}_{os.urandom(2).hex()}"
    work_dir.mkdir(parents=True, exist_ok=True)
    temp_path = work_dir / "work.docx"
    shutil.copy2(docx_path, temp_path)
    # 清除只读属性：copy2 会复制源文件的只读属性，导致临时文件无法删除
    try:
        os.chmod(temp_path, 0o666)
    except Exception:
        pass
    return str(temp_path), str(work_dir)


def _cleanup_temp(work_dir):
    """清理临时工作目录（多次重试以应对 Windows 文件锁）"""
    if not work_dir or not os.path.isdir(work_dir):
        return
    for attempt in range(5):
        try:
            # 先清除只读属性（copy2 复制的文件可能带只读属性，导致无法删除）
            for root, dirs, files in os.walk(work_dir):
                for f in files:
                    try:
                        os.chmod(os.path.join(root, f), 0o666)
                    except Exception:
                        pass
            shutil.rmtree(work_dir)
            return
        except Exception:
            if attempt < 4:
                time.sleep(0.2)
                gc.collect()
            else:
                pass  # 最后一次失败则忽略


def normalize_document_runs(docx_path):
    """
    将文档标准化：合并每段所有 run → 按句号分号拆分为独立句子 run。
    标准化后所有文档（不论来源）都具有一致的 run 结构，工具无需适配不同 run 分布。
    返回标准化后的文档路径。
    """
    import re
    temp_path, temp_dir = _copy_to_temp(docx_path)
    try:
        doc = Document(temp_path)
        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            # 1. 合并该段所有 run 文本
            full_text = "".join(r.text or "" for r in para.runs)

            # 2. 清空所有 run，只保留第一个
            runs_list = list(para.runs)
            first_run = runs_list[0]
            for run in runs_list[1:]:
                para._element.remove(run._element)

            # 更新第一个 run 的文本为合并后的全文
            t_el = first_run._element.find(qn('w:t'))
            if t_el is not None:
                t_el.text = full_text

            # 3. 按句子边界拆分（保留标点）
            sentences = re.split(r'(?<=[。；！？])', full_text)
            sentences = [s.strip() for s in sentences if s.strip()]
            if len(sentences) <= 1:
                continue

            # 获取第一个 run 的格式
            first_rPr = first_run._element.rPr

            # 第一个句子复用原 run
            if t_el is not None:
                t_el.text = sentences[0]

            # 后续句子各创建新 run
            for sent in sentences[1:]:
                new_r = OxmlElement("w:r")
                if first_rPr is not None:
                    new_r_rPr = deepcopy(first_rPr)
                    # 清理格式属性（由后续工具调用设置）
                    for tag in ("w:u", "w:highlight"):
                        old = new_r_rPr.find(qn(tag))
                        if old is not None:
                            new_r_rPr.remove(old)
                    new_r.append(new_r_rPr)
                t_el2 = OxmlElement("w:t")
                t_el2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                t_el2.text = sent
                new_r.append(t_el2)
                para._element.append(new_r)

        # 清除原文自带的黄色高亮（防止混淆工具判断）
        for para in doc.paragraphs:
            for run in para.runs:
                if run._element.rPr is not None:
                    hl_el = run._element.rPr.find(qn("w:highlight"))
                    if hl_el is not None and hl_el.get(qn("w:val")) != "none":
                        run._element.rPr.remove(hl_el)

        save_path = os.path.join(temp_dir, "normalized.docx")
        doc.save(save_path)
        del doc
        gc.collect()
        time.sleep(0.3)

        result_path = _safe_copy(save_path, docx_path)
        return result_path
    finally:
        _cleanup_temp(temp_dir)


def _safe_copy(src, dst):
    """
    安全复制：对原始作业文件，第一次生成 `{stem}_work.docx`。
    如果 dst 已经是 _work.docx，则直接覆盖（不再叠加 _work 后缀）。
    返回 dst 对应的 temp 路径。
    """
    _ensure_temp_dir()
    dst_path = Path(dst)
    stem = dst_path.stem
    suffix = dst_path.suffix

    # 如果已经是 _work 文件，直接覆盖
    if stem.endswith("_work"):
        new_dst = str(TEMP_DIR / f"{stem}{suffix}")
    else:
        new_dst = str(TEMP_DIR / f"{stem}_work{suffix}")

    shutil.copy2(src, new_dst)
    _temp_files.append(new_dst)
    return new_dst


def cleanup_temp_files(keep_paths: list = None):
    keep = set(keep_paths or [])
    for f in list(_temp_files):
        if f not in keep and os.path.exists(f):
            try:
                os.remove(f)
                _temp_files.remove(f)
            except Exception:
                pass


def _ensure_sentence_runs(para):
    """
    确保段落中每个句子（以。；！？结尾）都在独立的 run 中。
    如果某个 run 包含多个句子，将其拆分为多个 run。
    评语插入需要 sentence-level run 才能正确交错。
    """
    import re
    para_runs = list(para.runs)
    if len(para_runs) == 1 and len(para_runs[0].text or "") > 50:
        # 单个长 run → 按句子拆分
        text = para_runs[0].text or ""
        # 在句号、分号、问号、感叹号后拆分（保留标点）
        parts = re.split(r'(?<=[。；！？])', text)
        if len(parts) <= 1:
            return  # 只有一句，无需拆分

        first = True
        for part in parts:
            if not part.strip():
                continue
            if first:
                # 第一个部分复用原 run
                t_el = para_runs[0]._element.find(qn('w:t'))
                if t_el is not None:
                    t_el.text = part
                first = False
            else:
                # 后续部分创建新 run，复制原 run 的格式
                new_r = OxmlElement("w:r")
                if para_runs[0]._element.rPr is not None:
                    new_r_rPr = deepcopy(para_runs[0]._element.rPr)
                    # 清理格式属性（由后续工具调用设置）
                    for tag in ("w:u", "w:highlight"):
                        old = new_r_rPr.find(qn(tag))
                        if old is not None:
                            new_r_rPr.remove(old)
                    new_r.append(new_r_rPr)
                t_el = OxmlElement("w:t")
                # 保留 xml:space
                t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                t_el.text = part
                new_r.append(t_el)
                # 追加到段落末尾
                para._element.append(new_r)


# ============================================================
# 辅助：Run 拆分（确保高亮/下划线只覆盖精确的句子文本）
# ============================================================


def _split_run(run, split_char_index):
    """
    将一个 run 在指定字符位置拆分为两个 run。
    返回新创建的（后半段）run，若无需拆分则返回 None。
    """
    text = run.text or ""
    if split_char_index <= 0 or split_char_index >= len(text):
        return None

    # 创建新 run（后半段）
    new_r = OxmlElement("w:r")
    if run._element.rPr is not None:
        new_r_rPr = deepcopy(run._element.rPr)
        # 移除格式化属性，防止绿下划线/黄色高亮被复制到拆分后的 run 中
        for tag in ("w:u", "w:highlight"):
            old = new_r_rPr.find(qn(tag))
            if old is not None:
                new_r_rPr.remove(old)
        new_r.append(new_r_rPr)

    t_el = OxmlElement("w:t")
    t_el.text = text[split_char_index:]
    # 保留原 run 的 xml:space 属性
    old_t = run._element.find(qn('w:t'))
    if old_t is not None and old_t.get('{http://www.w3.org/XML/1998/namespace}space'):
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(t_el)

    # 截断原 run
    if old_t is not None:
        old_t.text = text[:split_char_index]

    # 在 run 后面插入新 run
    run._element.addnext(new_r)
    return new_r


def _trim_runs_to_sentence(para, sentence_start, sentence_end, target_runs):
    """
    将 target_runs 修剪到精确的 sentence 边界。
    如果某个 run 只有部分内容属于句子，将其拆分为两个 run。
    返回 (trimmed_runs, last_trimmed_run)
    """
    from docx.text.run import Run as DocxRun
    current_pos = 0
    trimmed_runs = []
    para_runs = list(para.runs)

    for run in para_runs:
        run_len = len(run.text or "")
        run_start = current_pos
        run_end = current_pos + run_len

        # 完全不重叠
        if run_end <= sentence_start or run_start >= sentence_end:
            current_pos += run_len
            continue

        # run 开始位置在句子前面 → 拆分
        if run_start < sentence_start:
            offset = sentence_start - run_start
            new_ct_r = _split_run(run, offset)
            # original run 现在是前缀（不在句子中），new_ct_r 是后缀（句子开头）
            current_pos += run_len  # 跳过整个原始 run（包括前缀和已拆分的后缀）
            run_start = sentence_start
            if new_ct_r is not None:
                new_run_len = len(new_ct_r.text or "")
                run_end = run_start + new_run_len
                # ⚠️ 新 run 可能仍超出 sentence_end，需要再次截断
                if run_end > sentence_end:
                    trim_offset = sentence_end - run_start
                    if 0 < trim_offset < new_run_len:
                        _split_run(DocxRun(new_ct_r, para._element), trim_offset)
                        # new_ct_r 已被截断到 sentence_end
                if run_start < sentence_end:
                    # 将 CT_R 包装为 Run 对象
                    new_run_obj = DocxRun(new_ct_r, para._element)
                    trimmed_runs.append(new_run_obj)
            continue  # 用分割后的新 run 替代，原 run 不加入

        # run 结束位置在句子后面 → 拆分
        if run_end > sentence_end:
            offset = sentence_end - run_start
            if 0 < offset < len(run.text or ""):
                _split_run(run, offset)
                # run 被截断到 sentence_end

        # 此 run 完全在句子内部
        trimmed_runs.append(run)
        current_pos += run_len

    return trimmed_runs, trimmed_runs[-1] if trimmed_runs else None


def _find_sentence_in_paragraph(para, sentence_text, trim=False):
    """
    在段落中定位句子文本所在的 run 列表。
    统一在归一化文本上匹配（兼容全角/半角、弯引号、零宽空格、重复标点、空白差异），
    再用 difflib 对齐把归一化位置映射回原文。
    如果 trim=True，使用 run 拆分功能确保精确边界（用于红色高亮）。
    返回 (target_runs, last_run) 或 None
    """
    import difflib
    para_text = para.text
    if not para_text:
        return None

    norm_para = _norm_text(para_text)
    norm_sent = _norm_text(sentence_text)
    if not norm_sent:
        return None

    # --- 1. 归一化精确匹配 ---
    idx_norm = norm_para.find(norm_sent)
    if idx_norm != -1:
        idx = _map_norm_index_to_original(para_text, norm_para, idx_norm)
        end = _map_norm_index_to_original(para_text, norm_para, idx_norm + len(norm_sent))
    else:
        # --- 2. 归一化模糊匹配（处理细微差异：重复标点、多余字符等） ---
        idx_norm, end_norm = _fuzzy_find_substring(norm_para, norm_sent)
        if idx_norm == -1:
            return None
        idx = _map_norm_index_to_original(para_text, norm_para, idx_norm)
        end = _map_norm_index_to_original(para_text, norm_para, end_norm)

    # 边界保护
    idx = max(0, min(idx, len(para_text)))
    end = max(idx, min(end, len(para_text)))

    # 根据 start/end 定位涉及的 runs
    current_pos = 0
    target_runs = []
    last_run = None
    for run in para.runs:
        run_len = len(run.text or "")
        run_start = current_pos
        run_end = current_pos + run_len
        if run_start < end and run_end > idx:
            target_runs.append(run)
            last_run = run
        current_pos += run_len

    if not target_runs:
        return None

    # 如需精确修剪（红色高亮用），拆分跨越边界的 run
    if trim:
        trimmed_runs, trimmed_last = _trim_runs_to_sentence(para, idx, end, target_runs)
        if trimmed_runs:
            return trimmed_runs, trimmed_last

    return target_runs, last_run


def _map_norm_index_to_original(original, norm_text, norm_idx):
    """
    将归一化后的索引映射回原文索引。
    支持长度变化（全角→半角、零宽空格删除、重复标点压缩等）。
    用 difflib 对齐原文与归一化文本，建立位置映射。
    """
    import difflib
    if norm_idx <= 0:
        return 0
    if norm_idx >= len(norm_text):
        return len(original)
    sm = difflib.SequenceMatcher(None, original, norm_text, autojunk=False)
    orig_pos = 0
    norm_pos = 0
    for o_start, n_start, size in sm.get_matching_blocks():
        # 匹配块之前的未匹配区：orig [orig_pos:o_start], norm [norm_pos:n_start]
        if norm_idx < n_start:
            # norm_idx 落在未匹配区 → 按比例映射回 orig
            unmatched_norm = norm_idx - norm_pos
            unmatched_orig = o_start - orig_pos
            if unmatched_norm <= 0:
                return orig_pos
            if unmatched_orig <= 0:
                return o_start
            # 按比例映射（未匹配区通常很小，近似线性）
            return orig_pos + int(unmatched_orig * unmatched_norm / max(1, n_start - norm_pos))
        # norm_idx 落在匹配块内
        if norm_idx < n_start + size:
            return o_start + (norm_idx - n_start)
        orig_pos = o_start + size
        norm_pos = n_start + size
    return len(original)


def _fuzzy_find_substring(text, query, threshold=0.6):
    """
    在 text 中滑动窗口搜索与 query 最相似的子串。
    返回 (start, end) 或 (-1, -1)。
    窗口差异范围较大，以适应重复标点、多余字符等细微差异。
    """
    import difflib
    if not query or not text:
        return -1, -1
    qlen = len(query)
    best_ratio = 0
    best_start = -1
    best_end = -1
    # 搜索窗口略大于 query，以适应标点差异（含重复标点、多余空格等）
    for win_delta in range(-6, 8):
        win_len = qlen + win_delta
        if win_len <= 0 or win_len > len(text):
            continue
        for i in range(len(text) - win_len + 1):
            candidate = text[i:i + win_len]
            ratio = difflib.SequenceMatcher(None, candidate, query).ratio()
            if ratio > best_ratio:
                best_ratio = ratio
                best_start = i
                best_end = i + win_len
    if best_ratio >= threshold:
        return best_start, best_end
    return -1, -1


def _to_chinese_punctuation(text: str) -> str:
    """将英文/半角标点替换为中文标点（用于评语内容），并去掉强调性单引号"""
    # 1. 去掉成对的弯单引号/直单引号（模型常用来强调词语，中文里形似英文撇号）
    text = re.sub(r"[\u2018\u2019']([^\u2018\u2019']{1,24})[\u2018\u2019']", r"\1", text)
    # 2. 成对直双引号 → 中文双引号
    text = re.sub(r'"([^"\n]{1,40})"', "“\\1”", text)
    # 3. 残余孤立引号
    text = text.replace('"', "“").replace("'", "’")
    # 4. 其余英文标点 → 中文
    text = text.replace(",", "，")
    text = text.replace("!", "！")
    text = text.replace("?", "？")
    text = text.replace(";", "；")
    text = text.replace(":", "：")
    text = text.replace("(", "（")
    text = text.replace(")", "）")
    text = text.replace("...", "…")
    text = text.replace("..", "…")
    text = text.replace("--", "——")
    text = text.replace("~", "～")
    return text


def _insert_comment_after_run(anchor_run, comment_text, color_rgb, newline=False):
    """在指定 run 的 XML 后面插入评语 run（自动转换英文标点为中文）。
    newline=True 时在评语前加换行符（<w:br/>），用于段落级评语单独成行。"""
    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color_rgb)
    rPr.append(color_el)

    new_r = OxmlElement("w:r")
    new_r.append(rPr)

    # 段落级评语：评语前加换行符，单独成行，避免与句子级评语挤在一起
    if newline:
        br_el = OxmlElement("w:br")
        new_r.append(br_el)

    t_el = OxmlElement("w:t")
    t_el.text = f"\u3010{_to_chinese_punctuation(comment_text)}\u3011"
    new_r.append(t_el)

    anchor_run._element.addnext(new_r)


# ============================================================
# 工具1: 读取参考答案 PDF
# ============================================================
def tool_read_answer_pdf(pdf_path: str) -> dict:
    """读取参考答案 PDF"""
    from pypdf import PdfReader
    reader = PdfReader(pdf_path)
    text = "\n".join(p.extract_text() for p in reader.pages)
    return {"text": text, "lines": len(text.splitlines())}


# ============================================================
# 工具2: 读取学生作业（句子级拆分）
# ============================================================
def tool_read_homework(docx_path: str) -> dict:
    """
    读取学生作业，提取【答案】部分，按句号和分号拆分为句子

    Returns:
        {
            "paragraphs": [...],
            "answer_start_index": 14,
            "sentences": [
                {"paragraph_idx": 0, "sentence_idx": 0, "text": "句子1。"},
                {"paragraph_idx": 0, "sentence_idx": 1, "text": "句子2。"},
                ...
            ],
            "sentence_count": 20
        }
    """
    text, paragraphs, start = extract_answer_part(docx_path)

    sentences = []
    for pi, para in enumerate(paragraphs):
        if not para.strip() or _is_protected_para(para):
            continue
        # 按句号、分号、问号、感叹号、换行拆分
        parts = re.split(r'(?<=[。；！？\n])', para)
        for si, part in enumerate(parts):
            part = part.strip().replace('\n', '')
            if part:
                sentences.append({
                    "paragraph_idx": pi,
                    "full_paragraph_idx": start + pi,  # 直接预计算全文索引
                    "sentence_idx": si,
                    "text": part
                })

    return {
        "paragraphs": paragraphs,
        "paragraph_count": len(paragraphs),
        "answer_start_index": start,
        "text": text,
        "sentences": sentences,
        "sentence_count": len(sentences),
    }


# ============================================================
# 工具4: 读取评分标准
# ============================================================
def tool_read_standard(standard_path: str) -> dict:
    """读取评分标准 Markdown"""
    with open(standard_path, "r", encoding="utf-8") as f:
        text = f.read()
    return {"text": text, "lines": len(text.splitlines())}


# ============================================================
# 文本定位助手（索引自修复）：模型不需要算准段落号，工具自动全文搜索
# ============================================================
def _is_comment_run(run):
    """判断 run 是否为本工具追加的评语（【开头或红/绿/蓝字色）"""
    text = (run.text or "").strip()
    if not text:
        return False
    if text.startswith("【"):
        return True
    rPr = run._element.rPr
    if rPr is not None:
        c = rPr.find(qn("w:color"))
        if c is not None and c.get(qn("w:val")) in ("00B050", "EE0000", "0070C0"):
            return True
    return False


def _para_clean_text(para):
    """段落中除去本工具评语后的原文文本"""
    return "".join((r.text or "") for r in para.runs if not _is_comment_run(r))


def _norm_text(t):
    """统一 Unicode：全角/半角、弯引号/直引号，用于文本匹配前的归一化"""
    t = t.replace("\u201c", "\"").replace("\u201d", "\"")   # 弯双引→直双引
    t = t.replace("\u2018", "'").replace("\u2019", "'")       # 弯单引→直单引
    t = t.replace("\uff0c", ",").replace("\u3000", " ")       # 全角逗号/全角空格
    t = t.replace("\uff01", "!").replace("\uff1f", "?")       # 全角感叹/问号
    t = t.replace("\uff1b", ";").replace("\uff1a", ":")       # 全角分号/冒号
    # 去掉不可见字符（零宽空格/零宽连接符/BOM 等），避免匹配失败
    for ch in ("\u200b", "\u200c", "\u200d", "\ufeff", "\u00a0"):
        t = t.replace(ch, "")
    # 压缩重复标点（模型常输出"。。""！！"等），统一为单个
    t = re.sub(r"([。！？；，,.;!?])\1+", r"\1", t)
    # 去掉所有空白字符（门控用宽松匹配，精确验证交给 _find_sentence_in_paragraph）
    t = "".join(ch for ch in t if not ch.isspace())
    return t


def _is_protected_para(text):
    """是否属于模板/非正文内容（题目、材料、说明行），禁止批改"""
    t = (text or "").strip()
    if not t:
        return True
    if t.startswith("【题目】") or t.startswith("【材料】") or t.startswith("【答案】"):
        return True
    if "示意批改讲师" in t or "可自行根据目标院校" in t:
        return True
    return False


def _find_sentence_across_paragraphs(doc, sentence_text, preferred_index=None, trim=False):
    """
    跨段落定位文本：优先在 preferred_index 段落查找，找不到则全文档搜索。
    只在“非评语文本”中匹配，避免匹配到已追加的评语内容。
    门控使用宽松匹配：先严格包含，失败则用相似度粗筛（放行给 _find_sentence_in_paragraph 精确判断）。
    返回 (paragraph, actual_index, (target_runs, last_run)) 或 (None, -1, None)
    """
    import difflib
    paras = doc.paragraphs
    order = []
    if preferred_index is not None and 0 <= preferred_index < len(paras):
        order.append(preferred_index)
    order += [i for i in range(len(paras)) if i not in order]
    n_sent = _norm_text(sentence_text)
    for idx in order:
        para = paras[idx]
        if not para.text.strip():
            continue
        n_para = _norm_text(_para_clean_text(para))
        if not n_para:
            continue
        # 门控：先严格包含（快速路径）
        if n_sent in n_para:
            result = _find_sentence_in_paragraph(para, sentence_text, trim=trim)
            if result:
                return para, idx, result
            continue
        # 门控放宽：相似度粗筛（句子核心部分在段落中，或整体相似度足够）
        # 用句子前 8 个字符做包含粗筛，避免长句误判
        sent_core = n_sent[:8]
        if sent_core and sent_core in n_para:
            result = _find_sentence_in_paragraph(para, sentence_text, trim=trim)
            if result:
                return para, idx, result
            continue
        # 最后兜底：整体相似度粗筛
        ratio = difflib.SequenceMatcher(None, n_sent, n_para).ratio()
        if ratio >= 0.5:
            result = _find_sentence_in_paragraph(para, sentence_text, trim=trim)
            if result:
                return para, idx, result
    return None, -1, None


def _locate_paragraph(doc, paragraph_index=None, anchor_text=None):
    """
    定位段落：优先用 anchor_text（该段包含的一个唯一短句）全文档搜索，其次用 paragraph_index。
    返回 (paragraph, actual_index) 或 (None, -1)
    """
    paras = doc.paragraphs
    if anchor_text:
        for i, p in enumerate(paras):
            if p.text.strip() and anchor_text in p.text:
                return p, i
    if paragraph_index is not None and 0 <= paragraph_index < len(paras) and paras[paragraph_index].text.strip():
        return paras[paragraph_index], paragraph_index
    return None, -1


# ============================================================
# 工具5: 句子级绿色评语（肯定）
# ============================================================
def tool_add_green_sentence(
    docx_path: str,
    sentence_text: str,
    comment: str,
    paragraph_index: int = None,
) -> dict:
    """
    对某个句子做肯定标记：原文加绿色下划线，句子后紧跟绿色评语【评语】。
    paragraph_index 只是提示，工具会在全文中自动定位 sentence_text。

    Args:
        docx_path: 当前文档路径
        sentence_text: 要批注的句子原文文本（必须精确匹配，含标点/引号）
        comment: 评语内容（不含【】）
        paragraph_index: 段落提示（可选，传错会自动纠正）
    """
    temp_path, temp_dir = _copy_to_temp(docx_path)
    try:
        doc = Document(temp_path)
        para, actual_idx, result = _find_sentence_across_paragraphs(
            doc, sentence_text, preferred_index=paragraph_index, trim=False
        )
        if not result:
            return {
                "docx_path": docx_path,
                "error": f"全文未找到文本: {sentence_text[:30]}。请确认与原文逐字一致（含标点/引号），可传更短、更精确的片段后重试。",
            }
        if _is_protected_para(para.text):
            return {"docx_path": docx_path, "skipped": True, "reason": "该段落是题目/材料/说明等非正文内容，跳过"}

        # 去重（在拆分前检查原始 run，避免标记被拆分剥离后漏检）
        for _r in result[0]:
            if _r._element.rPr is not None:
                _rPr = _r._element.rPr
                _u = _rPr.find(qn("w:u"))
                _hl = _rPr.find(qn("w:highlight"))
                if (_u is not None and _u.get(qn("w:val")) not in (None, "none")) or \
                   (_hl is not None and _hl.get(qn("w:val")) not in (None, "none")):
                    return {"docx_path": docx_path, "skipped": True, "reason": "该句子已有批注标记（下划线或高亮），跳过重复批注"}

        # 预处理：确保每个句子独立成 run，评语才能正确交错
        _ensure_sentence_runs(para)
        # 拆分后重新定位（run 结构变化，索引可能失效）
        result = _find_sentence_in_paragraph(para, sentence_text, trim=False)
        if not result:
            return {"docx_path": docx_path, "error": f"定位后仍未找到: {sentence_text[:30]}"}

        target_runs, last_run = result

        # 去重：该句子已有任何批注标记（下划线或高亮）→ 跳过，避免红绿混用/重复
        for run in target_runs:
            if run._element.rPr is not None:
                rPr = run._element.rPr
                u = rPr.find(qn("w:u"))
                hl = rPr.find(qn("w:highlight"))
                if (u is not None and u.get(qn("w:val")) not in (None, "none")) or \
                   (hl is not None and hl.get(qn("w:val")) not in (None, "none")):
                    return {"docx_path": docx_path, "skipped": True, "reason": "该句子已有批注标记（下划线或高亮），跳过重复批注"}

        # 原文加绿色下划线（跳过已有黄色高亮的 run，红优先）
        for run in target_runs:
            # 检查是否有黄色高亮
            has_hl = False
            if run._element.rPr is not None:
                hl_el = run._element.rPr.find(qn("w:highlight"))
                if hl_el is not None:
                    has_hl = True
            if has_hl:
                continue

            u_el = OxmlElement("w:u")
            u_el.set(qn("w:val"), "single")
            u_el.set(qn("w:color"), "00B050")
            if run._element.rPr is None:
                rPr = OxmlElement("w:rPr")
                run._element.insert(0, rPr)
            # 移除已有的 u 元素
            existing_u = run._element.rPr.find(qn("w:u"))
            if existing_u is not None:
                run._element.rPr.remove(existing_u)
            run._element.rPr.append(u_el)

        # 在句子后面插入绿色评语（评语为空时只加下划线不追加文字）
        if comment.strip():
            _insert_comment_after_run(last_run, comment, "00B050")

        save_path = os.path.join(temp_dir, "save.docx")
        doc.save(save_path)
        del doc
        gc.collect()
        time.sleep(0.3)

        result_path = _safe_copy(save_path, docx_path)
        return {"docx_path": result_path, "paragraph_index": actual_idx, "sentence": sentence_text[:30]}
    finally:
        _cleanup_temp(temp_dir)


# ============================================================
# 工具6: 句子级红色评语（纠错）
# ============================================================
def tool_add_red_sentence(
    docx_path: str,
    sentence_text: str,
    comment: str,
    paragraph_index: int = None,
) -> dict:
    """
    对某个文本片段做纠错标记：原文黄色高亮，句后紧跟红色评语【评语】。
    paragraph_index 只是提示，工具会在全文中自动定位 sentence_text。

    Args:
        docx_path: 当前文档路径
        sentence_text: 有问题的文本片段（将在此文本上加黄色高亮）
        comment: 评语内容（不含【】）
        paragraph_index: 段落提示（可选，传错会自动纠正）
    """
    temp_path, temp_dir = _copy_to_temp(docx_path)
    try:
        doc = Document(temp_path)
        para, actual_idx, result = _find_sentence_across_paragraphs(
            doc, sentence_text, preferred_index=paragraph_index, trim=True
        )
        if not result:
            return {
                "docx_path": docx_path,
                "error": f"全文未找到文本: {sentence_text[:30]}。请确认与原文逐字一致（含标点/引号），可传更短、更精确的片段后重试。",
            }
        if _is_protected_para(para.text):
            return {"docx_path": docx_path, "skipped": True, "reason": "该段落是题目/材料/说明等非正文内容，跳过"}

        # 去重（在拆分前检查原始 run，避免标记被拆分剥离后漏检）
        for _r in result[0]:
            if _r._element.rPr is not None:
                _rPr = _r._element.rPr
                _u = _rPr.find(qn("w:u"))
                _hl = _rPr.find(qn("w:highlight"))
                if (_u is not None and _u.get(qn("w:val")) not in (None, "none")) or \
                   (_hl is not None and _hl.get(qn("w:val")) not in (None, "none")):
                    return {"docx_path": docx_path, "skipped": True, "reason": "该片段已有批注标记（下划线或高亮），跳过重复批注"}

        # 预处理：确保每个句子独立成 run，评语才能正确交错
        _ensure_sentence_runs(para)
        # 拆分后重新定位（run 结构变化，索引可能失效）
        result = _find_sentence_in_paragraph(para, sentence_text, trim=True)
        if not result:
            return {"docx_path": docx_path, "error": f"定位后仍未找到: {sentence_text[:30]}"}

        target_runs, last_run = result

        # 去重：该片段已有任何批注标记（高亮或下划线）→ 跳过，避免红绿混用/重复
        for run in target_runs:
            if run._element.rPr is not None:
                rPr = run._element.rPr
                u = rPr.find(qn("w:u"))
                hl = rPr.find(qn("w:highlight"))
                if (u is not None and u.get(qn("w:val")) not in (None, "none")) or \
                   (hl is not None and hl.get(qn("w:val")) not in (None, "none")):
                    return {"docx_path": docx_path, "skipped": True, "reason": "该片段已有批注标记（下划线或高亮），跳过重复批注"}

        # 原文加黄色高亮（同时清除同一 run 上的绿色下划线，避免红绿混杂）
        for run in target_runs:
            # 清除绿色下划线（如果存在）
            if run._element.rPr is not None:
                existing_u = run._element.rPr.find(qn("w:u"))
                if existing_u is not None:
                    run._element.rPr.remove(existing_u)

            hl = OxmlElement("w:highlight")
            hl.set(qn("w:val"), "yellow")
            if run._element.rPr is None:
                rPr = OxmlElement("w:rPr")
                run._element.insert(0, rPr)
            run._element.rPr.append(hl)

        # 在句子后面插入红色评语
        _insert_comment_after_run(last_run, comment, "EE0000")

        save_path = os.path.join(temp_dir, "save.docx")
        doc.save(save_path)
        del doc
        gc.collect()
        time.sleep(0.3)

        result_path = _safe_copy(save_path, docx_path)
        return {"docx_path": result_path, "paragraph_index": actual_idx, "sentence": sentence_text[:30]}
    finally:
        _cleanup_temp(temp_dir)


# ============================================================
# 工具7: 段落级评语（整体诊断用）
# ============================================================
def tool_add_paragraph_comment(
    docx_path: str,
    comment: str,
    paragraph_index: int = None,
    color: str = "red",
    anchor_text: str = None,
) -> dict:
    """
    对整段追加评语（不匹配句子文本），用于整段的结构/逻辑/立意诊断。
    对应老师批改中的段落级评语，如"这段只有事实论证、没有观点收束"。
    可用 paragraph_index 或 anchor_text 定位段落。

    Args:
        docx_path: 当前文档路径
        comment: 评语内容（不含【】）
        paragraph_index: 段落索引（可选）
        color: "red"（问题，红色评语）或 "green"（肯定，绿色评语）
        anchor_text: 该段包含的一个唯一短句（可选，用于定位）
    """
    temp_path, temp_dir = _copy_to_temp(docx_path)
    try:
        doc = Document(temp_path)
        para, actual_idx = _locate_paragraph(doc, paragraph_index, anchor_text)
        if para is None:
            return {"docx_path": docx_path, "error": "未找到目标段落：请提供 paragraph_index 或 anchor_text（该段的一个唯一短句）"}
        if _is_protected_para(para.text):
            return {"docx_path": docx_path, "skipped": True, "reason": "该段落是题目/材料/说明等非正文内容，跳过"}

        color_hex = "EE0000" if color == "red" else "00B050"
        if para.runs:
            # 段落级评语换行单独成行，避免与句子级评语挤在一起
            _insert_comment_after_run(para.runs[-1], comment, color_hex, newline=True)
        else:
            run = para.add_run(f"【{_to_chinese_punctuation(comment)}】")
            run.font.color.rgb = COLOR_RED if color == "red" else COLOR_GREEN

        save_path = os.path.join(temp_dir, "save.docx")
        doc.save(save_path)
        del doc
        gc.collect()
        time.sleep(0.3)

        result_path = _safe_copy(save_path, docx_path)
        return {"docx_path": result_path, "paragraph_index": actual_idx, "comment": comment[:30], "color": color}
    finally:
        _cleanup_temp(temp_dir)


# ============================================================
# 工具8: 删除线标记（老师格式：标记可删内容）
# ============================================================
def tool_add_strikethrough(
    docx_path: str,
    sentence_text: str,
    paragraph_index: int = None,
) -> dict:
    """
    对原文片段加删除线（保留黑色字色），用于标记应删除/多余的内容
    （如冗余的"标题：""引语："标签、不必要的铺垫、多余年份等）。
    paragraph_index 只是提示，工具会在全文中自动定位。

    Args:
        docx_path: 当前文档路径
        sentence_text: 要加删除线的文本片段
        paragraph_index: 段落提示（可选）
    """
    temp_path, temp_dir = _copy_to_temp(docx_path)
    try:
        doc = Document(temp_path)
        para, actual_idx, result = _find_sentence_across_paragraphs(
            doc, sentence_text, preferred_index=paragraph_index, trim=True
        )
        if not result:
            return {
                "docx_path": docx_path,
                "error": f"全文未找到文本: {sentence_text[:30]}。请确认与原文逐字一致，可传更短、更精确的片段后重试。",
            }
        if _is_protected_para(para.text):
            return {"docx_path": docx_path, "skipped": True, "reason": "该段落是题目/材料/说明等非正文内容，跳过"}

        target_runs, _ = result
        for run in target_runs:
            if run._element.rPr is None:
                rPr = OxmlElement("w:rPr")
                run._element.insert(0, rPr)
            if run._element.rPr.find(qn("w:strike")) is None:
                strike = OxmlElement("w:strike")
                run._element.rPr.append(strike)

        save_path = os.path.join(temp_dir, "save.docx")
        doc.save(save_path)
        del doc
        gc.collect()
        time.sleep(0.3)

        result_path = _safe_copy(save_path, docx_path)
        return {"docx_path": result_path, "paragraph_index": actual_idx, "sentence": sentence_text[:30]}
    finally:
        _cleanup_temp(temp_dir)


# ============================================================
# 工具8.5: 表格单元格批注（采访提纲等含表格的作业）
# ============================================================
def tool_add_table_comment(
    docx_path: str,
    comment: str,
    table_index: int = 1,
    row_index: int = 0,
    col_index: int = 0,
) -> dict:
    """
    对表格中某个单元格添加 Word 批注（comment）。
    使用 Word 原生批注功能（w:commentRangeStart / w:commentRangeEnd / w:commentReference），
    批语显示在右侧批注栏，不改变单元格内文字。

    Args:
        docx_path: 当前文档路径
        comment: 评语内容
        table_index: 表格序号（从1开始）
        row_index: 行号（从0开始）
        col_index: 列号（从0开始）
    """
    temp_path, temp_dir = _copy_to_temp(docx_path)
    try:
        doc = Document(temp_path)
        if table_index < 1 or table_index > len(doc.tables):
            return {"docx_path": docx_path, "error": f"表格序号 {table_index} 超出范围（共 {len(doc.tables)} 个表格）"}
        table = doc.tables[table_index - 1]
        if row_index < 0 or row_index >= len(table.rows):
            return {"docx_path": docx_path, "error": f"行号 {row_index} 超出范围（共 {len(table.rows)} 行）"}
        row = table.rows[row_index]
        if col_index < 0 or col_index >= len(row.cells):
            return {"docx_path": docx_path, "error": f"列号 {col_index} 超出范围（共 {len(row.cells)} 列）"}
        cell = row.cells[col_index]
        if not cell.text.strip():
            return {"docx_path": docx_path, "skipped": True, "reason": "该单元格为空，跳过"}

        # 取单元格第一个段落的第一个 run 作为批注锚点
        para = cell.paragraphs[0] if cell.paragraphs else None
        if para is None or not para.runs:
            return {"docx_path": docx_path, "error": "单元格内无文本 run，无法添加批注"}

        anchor_run = para.runs[0]

        # 1. 用 python-docx 原生 API 创建批注内容（自动生成唯一 id + comments.xml 部件）
        comment_obj = doc.part.comments.add_comment(
            text=_to_chinese_punctuation(comment),
            author="xxxx",
            initials="SY",
        )
        comment_id = str(comment_obj.comment_id)

        # 2. 在锚点 run 前插入 commentRangeStart
        comment_range_start = OxmlElement("w:commentRangeStart")
        comment_range_start.set(qn("w:id"), comment_id)
        anchor_run._element.addprevious(comment_range_start)

        # 3. 在锚点 run 后插入 commentRangeEnd 和 commentReference
        comment_range_end = OxmlElement("w:commentRangeEnd")
        comment_range_end.set(qn("w:id"), comment_id)
        anchor_run._element.addnext(comment_range_end)

        comment_ref = OxmlElement("w:r")
        comment_ref_rPr = OxmlElement("w:rPr")
        comment_ref_rPr_style = OxmlElement("w:rStyle")
        comment_ref_rPr_style.set(qn("w:val"), "CommentReference")
        comment_ref_rPr.append(comment_ref_rPr_style)
        comment_ref.append(comment_ref_rPr)
        comment_ref_t = OxmlElement("w:commentReference")
        comment_ref_t.set(qn("w:id"), comment_id)
        comment_ref.append(comment_ref_t)
        comment_range_end.addnext(comment_ref)

        save_path = os.path.join(temp_dir, "save.docx")
        doc.save(save_path)
        del doc
        gc.collect()
        time.sleep(0.3)

        result_path = _safe_copy(save_path, docx_path)
        return {"docx_path": result_path, "table_index": table_index, "row_index": row_index, "col_index": col_index}
    finally:
        _cleanup_temp(temp_dir)


# ============================================================
# 工具9: 追加总评
# ============================================================
def tool_append_final_review(
    docx_path: str,
    score_breakdown: list,
    disclaimer: str = "",
    encouragement_paragraphs: list = None,
    signature: str = "",
) -> dict:
    """在文档末尾追加总评（全部蓝色 #0070C0，使用 add_run 防止中文拆 run）"""
    temp_path, temp_dir = _copy_to_temp(docx_path)
    try:
        doc = Document(temp_path)
        doc.add_paragraph("")

        def _add_blue_para(text):
            """添加一个蓝色段落，用单个 run 确保中文不分行，不设字号以继承原文默认"""
            p = doc.add_paragraph()
            run = p.add_run(_to_chinese_punctuation(text))
            run.font.color.rgb = COLOR_BLUE

        _add_blue_para("【总评】")

        for line in (score_breakdown or []):
            _add_blue_para(line)

        if disclaimer:
            _add_blue_para(disclaimer)

        # 合并鼓励段落：防止模型传成逐字列表
        if encouragement_paragraphs:
            # 过滤掉与签名重复的内容（模型常把"批改人：xxx"误放进鼓励段落）
            sig_norm = _norm_text(signature or "")
            merged = "".join(
                p for p in encouragement_paragraphs
                if not (sig_norm and _norm_text(p).strip() == sig_norm.strip())
            )
            # 按句号、分号、换行重新分段，不要逐字分段
            import re
            sentences = re.split(r'(?<=[。；！？])', merged)
            for sent in sentences:
                sent = sent.strip()
                if sent:
                    _add_blue_para(sent)

        if signature:
            _add_blue_para(signature)

        save_path = os.path.join(temp_dir, "save.docx")
        doc.save(save_path)
        del doc
        gc.collect()
        time.sleep(0.3)

        result_path = _safe_copy(save_path, docx_path)
        return {"docx_path": result_path}
    finally:
        _cleanup_temp(temp_dir)