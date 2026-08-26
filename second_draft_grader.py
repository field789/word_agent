"""
second_draft_grader.py —— 二稿批改器
============================================
核心理念：对比一稿批改的评语与二稿正文，判断学生是否修改，并插入反馈评语。

流程：
1. 从一稿批改 docx 提取评语（draft_common.extract_comments_from_graded）
2. 清洗二稿，提取答案正文段落（复用 single_shot_grader 的 LLM 清洗）
3. 调用 LLM：输入一稿评语列表 + 二稿答案正文，逐条判断每条评语是否已修改
4. 注入评语：已修改 → 绿色【✅】；未修改 → 红色评语（含"一稿已指出"提醒）
5. 追加蓝色总评（重新评分 + 对比一稿进步/退步）
6. 输出到二稿批改/complete_homework/

设计原则：
- 模块化：评语提取在 draft_common，LLM 调用复用 single_shot_grader._call_llm
- 评语注入复用 tools.py 的成熟工具（tool_add_green_sentence / tool_add_red_sentence）
- 评论二稿批改暂未实现（无范例），入口留空
"""

import json
import os
import re
import shutil
import sys
import time
from pathlib import Path
from typing import Optional

from docx import Document

# 确保项目根目录在路径中
_project_root = str(Path(__file__).resolve().parent.parent)
if _project_root not in sys.path:
    sys.path.insert(0, _project_root)

from word_agent.config import (
    ARK_API_KEY, DEFAULT_MODEL, DEFAULT_BASE_URL,
    TEMPERATURE, REASONING_EFFORT, MAX_TOKENS,
)
from word_agent.tools import (
    tool_add_green_sentence,
    tool_add_red_sentence,
    tool_add_paragraph_comment,
    tool_append_final_review,
    normalize_document_runs,
    cleanup_temp_files,
)
from word_agent.draft_common import (
    extract_comments_from_graded,
    extract_final_review,
    match_first_draft,
    normalize_stem,
)
from word_agent.single_shot_grader import _call_llm, _extract_answer_with_llm


# ============================================================
# 二稿批改评分维度（按作业类型）
# ============================================================
# 评论 50 分制 / 策划 30 分制 / 消息 40 分制
SECOND_DRAFT_DIMENSIONS = {
    "评论": {
        "score_total": 50,
        "dimensions": ["立意", "标题", "论证逻辑", "行文结构", "语言表达"],
    },
    "策划": {
        "score_total": 30,
        "dimensions": ["策划主题", "内容结构", "执行可行性"],
    },
    "消息": {
        "score_total": 40,
        "dimensions": ["标题", "导语", "逻辑清晰、事实清楚", "语言通顺、紧扣材料"],
    },
}


def _build_second_draft_system_prompt(homework_type: str) -> str:
    """按作业类型构建二稿批改 system prompt"""
    dims_cfg = SECOND_DRAFT_DIMENSIONS.get(homework_type, SECOND_DRAFT_DIMENSIONS["消息"])
    score_total = dims_cfg["score_total"]
    dims = dims_cfg["dimensions"]
    dims_str = "、".join(dims)
    breakdown_example = "、".join(f"{d}：X/10" for d in dims) + f"、总分：X/{score_total}"

    return f"""你是新闻传播学考研助教「时月学姐」，正在批改学生的{homework_type}作业**二稿**。
一稿已经批改过，现在你要对比一稿评语和二稿正文，判断学生是否针对一稿的评语做了修改，并给出反馈。

## 你的任务
1. 你会收到：一稿批改的评语列表（每条含：评语内容、对应的一稿原文）+ 二稿的答案正文
2. 对每条一稿评语，判断二稿中对应内容是否已修改
3. 对二稿中仍存在的问题（无论是否一稿指出过），给出红色评语
4. 最后输出二稿的总评（重新评分 + 对比一稿的进步/退步）

## 总评写作规范（必须遵守）
- 总评是**总结性的鼓励发言**，站在全文高度点评二稿整体表现，**绝不是对句评的重复或罗列**
- 禁止把句评内容搬进总评（如"你的标题改好了、第二段改对了、第三段还没改"这类逐条复述）
- 总评应聚焦：二稿整体相比一稿的**进步方向**（如结构更完整、观点更鲜明）、最值得肯定的整体亮点、仍需改进的**一个**核心方向（点到为止）
- 语气温暖向上，以鼓励收尾，让学生读完有方向、有动力
- 2-3 段即可，不要写成逐条点评的流水账

## 判断标准
- **已修改**：一稿评语指出的问题，在二稿中已改正 → 标记为已修改（绿色【✅】）
- **未修改**：一稿评语指出的问题，在二稿中仍存在 → 标记为未修改（红色评语，提醒"一稿已指出"）
- **部分修改**：部分改正但仍有问题 → 红色评语指出剩余问题
- **新问题**：二稿新出现的问题 → 红色评语指出

## 评语写作规范（必须遵守）
- 全中文标点：禁止任何半角英文标点（, . ? ! : ; " ' ( ) ~ 等），一律用中文标点
- 绿色评语用【✅】表示已修改到位，简洁肯定
- 红色评语直接给改法，必要时补一句原因；未修改的可引用一稿批语提醒
- **不要针对标点符号本身给出批改意见**：学生使用全角/半角标点、中英文标点混用等，一律不写纠错批语、不扣分、不点评。标点不是批改重点
- 称呼克制：不要每句都以"宝子"开头

## 评分维度（{score_total} 分制）
{dims_str}

## 输出格式（严格 JSON，不要任何额外文字）
```json
{{
  "inline_comments": [
    {{
      "type": "green",
      "paragraph_index": 3,
      "sentence_text": "二稿中已修改的原文句子（与二稿逐字一致）",
      "comment": "✅"
    }},
    {{
      "type": "red",
      "paragraph_index": 4,
      "sentence_text": "二稿中仍有问题的原文片段",
      "comment": "宝子，一稿已指出这个问题，仍未修改哦。正确应为..."
    }},
    {{
      "type": "paragraph_red",
      "paragraph_index": 5,
      "anchor_text": "该段唯一短句（可选）",
      "comment": "段落级问题评语"
    }}
  ],
  "final_review": {{
    "score_breakdown": ["{breakdown_example}"],
    "encouragement_paragraphs": ["总结性鼓励段落1（二稿整体进步点评，不重复句评）", "总结性鼓励段落2"],
    "signature": "批改人：时月学姐"
  }}
}}
```

## 技术要点
1. **paragraph_index 是全文段落索引**（即二稿用户 prompt 中 [段落 i] 的 i）
2. **green 类型**：sentence_text 传二稿中已修改的完整句子（含标点），与二稿逐字一致
3. **red 类型**：sentence_text 传二稿中仍有问题的原文片段（与二稿逐字一致），工具会精确高亮该片段
4. **paragraph_red/paragraph_green**：建议同时传 anchor_text
5. **每处一稿评语都要有反馈**：已修改的给 green，未修改的给 red，不能遗漏
6. **总评分数 = 各维度分数之和**
"""


# ============================================================
# 构建二稿批改的用户 prompt
# ============================================================

def _build_second_draft_user_prompt(
    first_draft_comments: list[dict],
    second_paragraphs: list[str],
    answer_indexes: list[int],
    first_draft_final_review: list[str],
) -> str:
    """构建二稿批改的用户 prompt"""
    # 一稿评语列表
    comment_lines = []
    for i, c in enumerate(first_draft_comments):
        orig = c.get("original_text", "")[:60]
        comment_lines.append(
            f"[评语{i}] 颜色={c.get('color')} | 对应原文: {orig} | 评语: {c.get('comment')}"
        )
    first_comments_text = "\n".join(comment_lines) if comment_lines else "（无评语）"

    # 一稿总评
    first_review_text = "\n".join(first_draft_final_review) if first_draft_final_review else "（无总评）"

    # 二稿答案正文（带全文索引）
    from word_agent.tools import _is_protected_para
    para_lines = []
    for i, p in enumerate(second_paragraphs):
        if not p.strip() or i not in answer_indexes:
            continue
        # 防御：过滤题目/材料/说明等保护段落（清洗可能偶发漏排）
        if _is_protected_para(p):
            continue
        para_lines.append(f"[段落 {i}] {p}")
    second_text = "\n\n".join(para_lines)

    return (
        f"## 一稿批改评语\n{first_comments_text}\n\n"
        f"## 一稿总评\n{first_review_text}\n\n"
        f"## 二稿答案正文（带全文段落索引）\n{second_text}\n\n"
        "请对比一稿评语与二稿正文，逐条判断是否已修改，并输出二稿批改 JSON。"
    )


# ============================================================
# 二稿批改主流程
# ============================================================

def run_second_draft_grader(
    first_draft_docx_path: str,
    second_draft_docx_path: str,
    homework_type: str = "消息",
    output_docx_path: Optional[str] = None,
    author: str = "时月学姐",
) -> dict:
    """
    运行二稿批改流程

    Args:
        first_draft_docx_path: 一稿批改 docx 路径（含评语）
        second_draft_docx_path: 二稿原始 docx 路径
        homework_type: 作业类型（"消息" / "评论"）
        output_docx_path: 输出路径（可选，默认自动生成）
        author: 批改人署名

    Returns:
        dict: 批改结果汇总
    """
    SEP = "=" * 60
    print(SEP)
    print(f"Word Agent 二稿批改系统 ({homework_type} 二稿)")
    print(SEP)

    # ---- 第1步：提取一稿评语 ----
    print("\n📖 第1步：提取一稿批改评语...")
    first_comments = extract_comments_from_graded(first_draft_docx_path)
    first_review = extract_final_review(first_draft_docx_path)
    print(f"   ✅ 一稿评语: {len(first_comments)} 条")
    print(f"   ✅ 一稿总评: {len(first_review)} 段")

    # ---- 第2步：标准化二稿文档 ----
    print("\n🔧 第2步：标准化二稿文档 run 结构...")
    current_docx = normalize_document_runs(second_draft_docx_path)
    print(f"   ✅ 已标准化")

    # ---- 第3步：清洗二稿，提取答案正文 ----
    print("\n🧹 第3步：LLM 清洗二稿，识别答案正文...")
    answer_indexes, full_paragraphs = _extract_answer_with_llm(second_draft_docx_path)
    if not answer_indexes:
        print("   ⚠️  清洗未识别到答案正文，回退为全文批改")
        answer_indexes = list(range(len(full_paragraphs)))
    print(f"   ✅ 答案正文段落: {answer_indexes}")

    # ---- 第4步：调用 LLM 对比判断 ----
    print("\n🤖 第4步：调用 DeepSeek 对比一稿评语与二稿...")
    print("   正在分析，请稍候...")
    system_prompt = _build_second_draft_system_prompt(homework_type)
    user_prompt = _build_second_draft_user_prompt(
        first_comments,
        full_paragraphs,
        answer_indexes,
        first_review,
    )
    result = _call_llm(system_prompt, user_prompt)

    # ---- 第5步：注入评语 ----
    print("\n✏️  第5步：注入二稿评语...")
    comments = result.get("inline_comments", [])
    print(f"   共 {len(comments)} 条评语")
    current_docx = _inject_second_draft_comments(current_docx, comments)

    # ---- 第6步：追加总评 ----
    print("\n📝 第6步：追加总评...")
    final_review = result.get("final_review", {})
    if final_review and final_review.get("score_breakdown"):
        fr_result = tool_append_final_review(
            docx_path=current_docx,
            score_breakdown=final_review.get("score_breakdown", []),
            disclaimer=final_review.get("disclaimer", ""),
            encouragement_paragraphs=final_review.get("encouragement_paragraphs", []),
            signature=final_review.get("signature", "") or f"批改人：{author}",
        )
        if "docx_path" in fr_result:
            current_docx = fr_result["docx_path"]
        print("   ✅ 总评已追加")
    else:
        print("   ⚠️  未生成总评（JSON 中缺少 final_review.score_breakdown）")

    # ---- 第7步：输出 ----
    print("\n💾 第7步：保存文件...")
    if output_docx_path is None:
        src = Path(second_draft_docx_path)
        # 当日日期子文件夹（如 2026-8-19），输出到 二稿批改/二稿批改结果/{日期}/
        now = time.localtime()
        date_dir = f"{now.tm_year}-{now.tm_mon}-{now.tm_mday}"
        output_docx_path = str(src.parent.parent / "二稿批改结果" / date_dir / f"【已批改】{src.stem}{src.suffix}")

    os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
    try:
        shutil.copy2(current_docx, output_docx_path)
        print(f"   ✅ 输出文件: {output_docx_path}")
    except Exception as e:
        print(f"   ❌ 复制失败: {e}")
        print(f"   临时文件保留在: {current_docx}")

    cleanup_temp_files(keep_paths=[current_docx])

    print("\n" + SEP)
    print("✅ 二稿批改完成！")
    print(SEP)

    return {
        "output_path": output_docx_path,
        "comments_count": len(comments),
        "first_comments_count": len(first_comments),
    }


# ============================================================
# 二稿评语注入（复用 tools.py 工具）
# ============================================================

def _inject_second_draft_comments(current_docx: str, comments: list) -> str:
    """逐条注入二稿评语到文档"""
    if not comments:
        print("   📭 没有需要批注的评论")
        return current_docx

    ok = 0
    failed = 0
    skipped = 0

    for i, c in enumerate(comments):
        ctype = c.get("type", "")
        para_idx = c.get("paragraph_index", 0)
        comment_text = c.get("comment", "")
        sentence_text = c.get("sentence_text", "")
        anchor_text = c.get("anchor_text", "")

        try:
            if ctype == "paragraph_red":
                result = tool_add_paragraph_comment(
                    docx_path=current_docx,
                    comment=comment_text,
                    paragraph_index=para_idx,
                    color="red",
                    anchor_text=anchor_text or None,
                )
            elif ctype == "paragraph_green":
                result = tool_add_paragraph_comment(
                    docx_path=current_docx,
                    comment=comment_text,
                    paragraph_index=para_idx,
                    color="green",
                    anchor_text=anchor_text or None,
                )
            elif ctype == "green":
                if not sentence_text:
                    print(f"   ⚠️  [{i}] green 类型缺少 sentence_text，跳过")
                    failed += 1
                    continue
                result = tool_add_green_sentence(
                    docx_path=current_docx,
                    sentence_text=sentence_text,
                    comment=comment_text,
                    paragraph_index=para_idx,
                )
            elif ctype == "red":
                if not sentence_text:
                    print(f"   ⚠️  [{i}] red 类型缺少 sentence_text，跳过")
                    failed += 1
                    continue
                result = tool_add_red_sentence(
                    docx_path=current_docx,
                    sentence_text=sentence_text,
                    comment=comment_text,
                    paragraph_index=para_idx,
                )
            else:
                print(f"   ⚠️  [{i}] 未知类型: {ctype}，跳过")
                failed += 1
                continue

            if "docx_path" in result:
                current_docx = result["docx_path"]

            if result.get("error"):
                print(f"   ❌ [{i}] {ctype}: {result['error'][:60]}")
                failed += 1
            elif result.get("skipped"):
                print(f"   ⏭️  [{i}] {ctype} 段落{para_idx}: {result.get('reason', '跳过')[:40]}")
                skipped += 1
            else:
                label = {"green": "🟢", "red": "🔴", "paragraph_red": "🔴段", "paragraph_green": "🟢段"}.get(ctype, ctype)
                print(f"   ✅ [{i}] {label} 段落{para_idx}")
                ok += 1

        except Exception as e:
            print(f"   ❌ [{i}] {ctype} 异常: {e}")
            failed += 1

    print(f"\n   📊 二稿批注注入结果: ✅ {ok} 成功 / ⏭️  {skipped} 跳过 / ❌ {failed} 失败 / 共 {len(comments)} 条")
    return current_docx
