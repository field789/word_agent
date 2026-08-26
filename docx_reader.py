"""
docx_reader.py —— 容错 docx 读取模块
============================================
问题背景：
  某些 docx 文件的 zip 中央目录 CRC-32 元数据损坏（如 word/media/image1.png），
  导致 python-docx / zipfile 读取时报 BadZipFile 错误，但 Word 能正常打开
  （Word 容错强，忽略损坏的图片；python-docx 严格校验，遇到即报错）。

方案：
  重新打包 docx：绕过损坏的 CRC 元数据，读取每个条目的原始字节，
  重新计算正确的 CRC-32 写入新 zip。图片数据本身通常是完整的，只是 CRC 值错了。

用法：
  from word_agent.docx_reader import open_docx_safe
  doc = open_docx_safe(path)   # 自动修复损坏文件后返回 Document
"""

import os
import struct
import tempfile
import zipfile
import zlib

from docx import Document


def _read_raw_entry(zip_path: str, info) -> bytes:
    """绕过 CRC 校验，读取 zip 条目的原始字节"""
    with open(zip_path, 'rb') as f:
        f.seek(info.header_offset)
        header = f.read(30)
        name_len, extra_len = struct.unpack('<HH', header[26:30])
        data_start = info.header_offset + 30 + name_len + extra_len
        f.seek(data_start)
        raw = f.read(info.compress_size)
    if info.compress_type == 0:  # stored（不压缩）
        return raw
    return zlib.decompress(raw, -15)  # deflated


def repair_docx(src_path: str, dst_path: str = None) -> str:
    """
    重新打包 docx，修复 CRC 错误。
    读取每个条目的原始字节（绕过损坏的 CRC 元数据），重新计算正确 CRC 写入新 zip。
    返回修复后的路径。

    注意：支持 src_path == dst_path（原地修复）。此时会先把所有条目读入内存，
    关闭输入文件后再写目标文件，避免读写同一文件导致文件被截断损坏。
    """
    if dst_path is None:
        # 在源文件同目录创建临时文件，避免跨盘移动问题
        src_dir = os.path.dirname(os.path.abspath(src_path))
        fd, dst_path = tempfile.mkstemp(suffix='.docx', dir=src_dir)
        os.close(fd)

    in_place = os.path.abspath(src_path) == os.path.abspath(dst_path)

    zin = zipfile.ZipFile(src_path)
    try:
        # 先把所有条目读入内存（绕过损坏的 CRC 元数据）
        entries = []
        for info in zin.infolist():
            data = _read_raw_entry(src_path, info)
            entries.append((info.filename, data))
    finally:
        zin.close()

    # 原地修复时，输入文件已关闭，此时再写目标文件不会互相干扰
    zout = zipfile.ZipFile(dst_path, 'w', zipfile.ZIP_DEFLATED)
    try:
        for name, data in entries:
            zout.writestr(name, data)
    finally:
        zout.close()
    return dst_path


def _clear_readonly(path: str):
    """清除 Windows 文件的只读属性（只读文件无法被 os.replace/os.remove）"""
    try:
        import ctypes
        FILE_ATTRIBUTE_READONLY = 1
        attrs = os.stat(path).st_file_attributes
        if attrs & FILE_ATTRIBUTE_READONLY:
            ctypes.windll.kernel32.SetFileAttributesW(path, attrs & ~FILE_ATTRIBUTE_READONLY)
    except Exception:
        pass


def open_docx_safe(docx_path: str) -> Document:
    """
    安全打开 docx：先用 python-docx 打开，失败（CRC 损坏）则修复后重试。
    返回 Document 对象。

    注意：修复后的文件会替换原文件（清除只读属性后 os.replace）。
    若替换失败（Windows 文件锁），则返回修复后临时文件的 Document，
    调用方需用 doc.save() 保存到目标路径。
    """
    try:
        return Document(docx_path)
    except (zipfile.BadZipFile, Exception) as e:
        # 尝试修复（重新计算 CRC）
        fixed_path = None
        try:
            fixed_path = repair_docx(docx_path)
            # 验证修复后的文件可打开（打开后立即关闭，释放文件句柄）
            doc = Document(fixed_path)
            del doc
            import gc
            gc.collect()
            # 尝试替换原文件（先清除只读属性）
            try:
                _clear_readonly(docx_path)
                os.replace(fixed_path, docx_path)
                return Document(docx_path)
            except Exception:
                # 替换失败（文件被占用等）：直接返回修复后临时文件的 Document
                return Document(fixed_path)
        except Exception as repair_err:
            # 清理可能残留的临时文件
            try:
                if fixed_path and os.path.exists(fixed_path):
                    os.remove(fixed_path)
            except Exception:
                pass
            raise RuntimeError(
                f"无法读取 docx: {docx_path}\n"
                f"原始错误: {e}\n"
                f"修复失败: {repair_err}"
            ) from e
